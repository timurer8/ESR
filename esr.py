    # Версия v1.8.5-beta
    # 📦 Блок 1: Импорты, Инициализация и Конструктор GUI
import sys, os, re, asyncio, threading, json, customtkinter as ctk, edge_tts, io
from tkinter import filedialog, messagebox
from bs4 import BeautifulSoup
os.environ['PYGAME_HIDE_SUPPORT_PROMPT'] = "hide"
import pygame

# Инициализируем базовую тему приложения
ctk.set_appearance_mode("Light")
ctk.set_default_color_theme("blue")

class EdgeStreamReaderApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("ESR v1.8.5-beta - EdgeStream Reader [Релиз]")
        self.geometry("720x710")
        self.resizable(False, False)
        pygame.mixer.init()
        
        self.config_file = "esr_config.json"
        self.bookmarks_file = "esr_bookmarks.json"
        self.last_file_path = ""
        self.settings_window = None  # Ссылка на активное дочернее окно настроек
        
        self.theme = "Светлая (Light)"
        self.text_bg = "«Стандартный»"
        
        # Карта доступных голосов (Онлайн Edge-TTS и Офлайн системные SAPI5)
        self.voice_map = {
            "Дмитрий (Онлайн Edge)": "ru-RU-DmitryNeural",
            "Светлана (Онлайн Edge)": "ru-RU-SvetlanaNeural",
            "Дарья (Офлайн SAPI5)": "TTS_MS_RU-RU_DARIYA_NATIVE",
            "Дмитрий (Офлайн SAPI5)": "TTS_MS_RU-RU_DMITRY_NATIVE",
            "Екатерина (Офлайн SAPI5)": "TTS_MS_RU-RU_EKATERINA_NATIVE",
            "Светлана (Офлайн SAPI5)": "TTS_MS_RU-RU_SVETLANA_NATIVE"
        }
        self.voice, self.speed, self.volume = "ru-RU-DmitryNeural", "0", "0"
        
        self.is_speaking = False
        self.is_paused = False
        self.paragraphs = []         # Исходные авторские абзацы текста для отображения в GUI
        self.virtual_chunks = []      # Структурированные виртуальные чанки: [(para_idx, chunk_text), ...]
        self.current_chunk_index = 0# Глобальный указатель (маркер) текущего чанка чтения
        self.async_loop = None
        # RAM-буфер: Хранилище аудио-потоков MP3/WAV в оперативной памяти без обращений к диску
        self.ram_cache = {} # Хранилище MP3-чанков в ОЗУ [chunk_idx: io.BytesIO]
        
        # --- ИНИЦИАЛИЗАЦИЯ ДВИЖКА СЛОВАРЕЙ DEMAGOG v1.8.5 ---
        self.exact_replacements = {}  # Сюда будут динамически собираться активные точные правила
        self.regex_replacements = []  # Сюда будут собираться активные маски (*)
        self.all_dicts_database = {}
        
        # Глобальное хранилище ВСЕХ правил, разбитое по именам файлов
        self.all_dicts_database = {}  
        
        # Состояние чекбоксов в GUI (по умолчанию ВСЕ ВЫКЛЮЧЕНЫ = False)
        self.dict_states = {
            "Edge_new": False,
            "edge_Yo_": False,
            "My_EDGE_new": False,
            "RU_EDGE_new": False
        }

        # --- БЛОК ЛОКАЛИЗАЦИИ ИНТЕРФЕЙСА v1.8.5 ---
        self.lang = "RU"  # По умолчанию русский язык
        self.translations = {
            "RU": {
                "title": "ESR v1.8.5-beta - EdgeStream Reader",
                "btn_load": " 📁 Открыть книгу (.txt, .fb2)",
                "btn_settings": " ⚙ Настройки",
                "btn_bookmark": " + 🔖 Закладка",
                "btn_view_bookmarks": " 📚 Закладки",
                "lbl_file_empty": "Файл не выбран",
                "text_area_init": "Вставьте текст или откройте файл книги...",
                "speed_label": "Скорость темпа: ",
                "btn_reader_mode": " 📖 Окно чтения",
                "btn_play": "► Озвучить",
                "btn_pause": " ⏸ Пауза",
                "btn_resume": "► Продолжить",
                "btn_stop": " ⬛ Стоп",
                "btn_export": " 💾🎵 Сохранить Аудио",
                "menu_cut": " ✂ Вырезать",
                "menu_copy": " 📋 Копировать",
                "menu_paste": " 📥 Вставить",
                "tab_interface": "Интерфейс",
                "tab_help": "Помощь",
                "tab_about": "О программе",
                "lbl_theme": "Тема оформления приложения:",
                "lbl_bg": "Фон окна читаемого текста:",
                "lbl_lang": "Язык интерфейса (Language):",
                "btn_open_dict": " 🗂 Открыть пользовательский словарь",
                "reader_title": "ESR v1.8.5-beta — Экран чтения",
                "reader_bookmark": " 📍 Поставить закладку",
                "progress_read": "Чтение чанков: ",
                "progress_pause": "Пауза: ",
                "progress_done": "Завершено! 100%",
                "dict_win_title": "Пользовательский словарь ESR",
                "dict_empty": "Словарь пуст. Добавьте свои правила в формате слово=замена.",
                "dict_placeholder": "формат: слово=замена",
                "dict_btn_add": " + Добавить",
                "dict_selected": "Выбрано слово: ",
                "dict_none": "[нет]",
                "dict_btn_del": " ❌ Удалить",
                "lbl_demagog": "Словари Demagog:",
                "bg_standard": "«Стандартный»",
                "bg_sepia": "«Книжная Сепия»",
                "bg_oled": "«Ночной OLED»",
                "bm_title": " 📚 Список закладок",
                "bm_info": "Двойной клик — ПЕРЕХОД. Одиночный клик + Кнопка — УДАЛЕНИЕ.",
                "bm_btn_del": " ❌ Удалить выбранную закладку",
                "bm_pos": "Поз.",
                "msg_success": "Успех",
                "msg_error": "Ошибка",
                "msg_warning": "Внимание",
                "msg_no_book": "Книга не выбрана! Сначала откройте файл книги.",
                "msg_bm_added": "Закладка успешно добавлена на позиции: ",
                "msg_bm_deleted": "Закладка успешно удалена.",
                "msg_bm_not_found": "Не удалось найти или удалить выбранную закладку.",
                "msg_audio_saved": "Аудио успешно экспортировано и сохранено в файл!",
                "msg_dict_saved": "Пользовательский словарь успешно обновлен и сохранен.",
                "msg_empty_fields": "Поля ввода не могут быть пустыми!"
            },
            "EN": {
                "title": "ESR v1.8.5-beta - EdgeStream Reader",
                "btn_load": " 📁 Open Book (.txt, .fb2)",
                "btn_settings": " ⚙ Settings",
                "btn_bookmark": " + 🔖 Bookmark",
                "btn_view_bookmarks": " 📚 Bookmarks",
                "lbl_file_empty": "File not selected",
                "text_area_init": "Paste text or open a book file...",
                "speed_label": "Speech Rate: ",
                "btn_reader_mode": " 📖 Reader Mode",
                "btn_play": "► Play Speech",
                "btn_pause": " ⏸ Pause",
                "btn_resume": "► Resume",
                "btn_stop": " ⬛ Stop",
                "btn_export": " 💾🎵 Save Audio",
                "menu_cut": " ✂ Cut",
                "menu_copy": " 📋 Copy",
                "menu_paste": " 📥 Paste",
                "tab_interface": "Interface",
                "tab_help": "Help",
                "tab_about": "About",
                "lbl_theme": "Application Appearance Theme:",
                "lbl_bg": "Text Window Background Color:",
                "lbl_lang": "Interface Language:",
                "btn_open_dict": " 🗂 Open User Pronunciation Dictionary",
                "reader_title": "ESR v1.8.5-beta — Reader Screen",
                "reader_bookmark": " 📍 Place Bookmark",
                "progress_read": "Reading chunks: ",
                "progress_pause": "Paused: ",
                "progress_done": "Finished! 100%",
                "dict_win_title": "ESR User Pronunciation Dictionary",
                "dict_empty": "Dictionary is empty. Add your rules in word=replacement format.",
                "dict_placeholder": "format: word=replacement",
                "dict_btn_add": " + Add Rule",
                "dict_selected": "Selected word: ",
                "dict_none": "[none]",
                "dict_btn_del": " ❌ Delete",
                "lbl_demagog": "Demagog Dictionaries:",
                "bg_standard": "Default Style",
                "bg_sepia": "Book Sepia",
                "bg_oled": "Night OLED",
                "bm_title": " 📚 Bookmark List",
                "bm_info": "Double click — GO TO. Single click + Button — DELETE.",
                "bm_btn_del": " ❌ Delete Selected Bookmark",
                "bm_pos": "Pos.",
                "msg_success": "Success",
                "msg_error": "Error",
                "msg_warning": "Warning",
                "msg_no_book": "No book selected! Please open a book file first.",
                "msg_bm_added": "Bookmark successfully added at position: ",
                "msg_bm_deleted": "Bookmark successfully deleted.",
                "msg_bm_not_found": "Failed to find or delete the selected bookmark.",
                "msg_audio_saved": "Audio successfully exported and saved to file!",
                "msg_dict_saved": "User dictionary successfully updated and saved.",
                "msg_empty_fields": "Input fields cannot be empty!"
            }
        }

        # Первичный разбор файлов с диска в память (без активации в плеере)
        self.load_demagog_dictionaries()
        # ----------------------------------------------------
        
        self.prefetch_thread = None
        self.is_prefetching = False
        self.is_exporting = False
        
        self.create_widgets()
        self.load_state() 
        
        if os.path.exists("icon.ico"):
            try:
                self.wm_iconbitmap("icon.ico")
            except: pass
            
        self.protocol("WM_DELETE_WINDOW", self.on_close)

    def tr(self, key):
        """Безопасный метод получения перевода по ключу"""
        return self.translations.get(self.lang, self.translations["RU"]).get(key, key)

    def create_widgets(self):
        self.file_frame = ctk.CTkFrame(self)
        self.file_frame.pack(fill="x", padx=20, pady=15)
        
        self.btn_load = ctk.CTkButton(self.file_frame, text="📁 Открыть книгу (.txt, .fb2)", command=self.load_file, width=180)
        self.btn_load.pack(side="left", padx=5, pady=5)
        
        self.btn_settings = ctk.CTkButton(self.file_frame, text="⚙ Настройки", command=self.open_settings, width=100, fg_color="#4b5563", hover_color="#374151")
        self.btn_settings.pack(side="left", padx=5, pady=5)
        
        self.btn_add_bookmark = ctk.CTkButton(self.file_frame, text=" 🔖 +Закладка", command=self.add_bookmark, width=100, fg_color="#0284c7", hover_color="#0369a1")
        self.btn_add_bookmark.pack(side="left", padx=5, pady=5)
        
        self.btn_view_bookmarks = ctk.CTkButton(self.file_frame, text=" 📚 Закладки", command=self.open_bookmarks, width=100, fg_color="#4f46e5", hover_color="#4338ca")
        self.btn_view_bookmarks.pack(side="left", padx=5, pady=5)
        
        self.lbl_file = ctk.CTkLabel(self.file_frame, text="Файл не выбран", text_color="gray", font=("Arial", 13))
        self.lbl_file.pack(side="left", padx=10, pady=5)
        
        self.text_area = ctk.CTkTextbox(self, width=680, height=340, font=("Arial", 14),
            text_color=("black", "white"), fg_color=("white", "#1f2937"),
            exportselection=False)
        self.text_area.pack(pady=10, padx=20)
        self.text_area.insert("1.0", "Вставьте текст или откройте файл книги...")
        
        self.text_area.bind("<Button-1>", self._on_text_click)
        
        self.text_area.tag_config("para_highlight", background="#ccfbf1", foreground="black")
        
        self.progress_frame = ctk.CTkFrame(self)
        self.progress_frame.pack(fill="x", padx=20, pady=5)
        self.progress_label = ctk.CTkLabel(self.progress_frame, text="Прогресс: 0%", font=("Arial", 12, "bold"))
        self.progress_label.pack(side="left", padx=10)
        self.progress_bar = ctk.CTkProgressBar(self.progress_frame, orientation="horizontal", height=12)
        self.progress_bar.pack(side="left", fill="x", expand=True, padx=10)
        self.progress_bar.set(0)

        self.controls_frame = ctk.CTkFrame(self)
        self.controls_frame.pack(fill="x", padx=20, pady=10)
        
        self.speed_label = ctk.CTkLabel(self.controls_frame, text="Скорость темпа: 0", font=("Arial", 13))
        self.speed_label.pack(side="left", padx=15)
        
        # Создаем ползунок строго внутри self.controls_frame, а не в self!
        self.speed_slider = ctk.CTkSlider(self.controls_frame, from_=-10, to=10, number_of_steps=20, command=self.on_speed_change)
        self.speed_slider.set(0)
        self.speed_slider.pack(side="left", padx=10, expand=True, fill="x")
        
        self.voice_combo = ctk.CTkComboBox(self.controls_frame, values=list(self.voice_map.keys()), command=self.change_voice, width=180)
        self.voice_combo.set("Дмитрий (Мужской)")
        self.voice_combo.pack(side="right", padx=15, pady=5)

        # --- БЛОК ЧЕКБОКСОВ ДЛЯ СЛОВАРЕЙ DEMAGOG v1.8.5 ---
        dict_frame = ctk.CTkFrame(self, fg_color="transparent")
        dict_frame.pack(pady=5, padx=20, fill="x")
        
        # Добавляем self., чтобы метка была доступна для перевода из любого места программы
        self.lbl_demagog_title = ctk.CTkLabel(dict_frame, text="Словари Demagog:", font=("Arial", 11, "bold"))
        self.lbl_demagog_title.pack(side="left", padx=5)

        # Функция-обработчик клика по флажку
        def toggle_dict(dict_name, var):
            self.dict_states[dict_name] = bool(var.get())
            self.rebuild_active_replacements()

        # Создаем чекбокс для каждого словаря из нашей базы данных
        for dict_name in sorted(self.dict_states.keys()):
            var = ctk.IntVar(value=0) # По умолчанию выключен (0)
            cb = ctk.CTkCheckBox(dict_frame, text=dict_name, variable=var,
                                 font=("Arial", 10), checkbox_width=16, checkbox_height=16,
                                 command=lambda name=dict_name, v=var: toggle_dict(name, v))
            cb.pack(side="left", padx=8)
        # --------------------------------------------------

        # --- ДОРАБОТКА (Дизайн кнопок): ВЫРАЗИТЕЛЬНЫЕ СИСТЕМНЫЕ СИМВОЛЫ МЕДИА-ПАНЕЛИ ---
        self.buttons_frame = ctk.CTkFrame(self)
        self.buttons_frame.pack(fill="x", padx=20, pady=15)

        self.btn_reader_mode = ctk.CTkButton(self, text="📖 Окно чтения", command=self.open_reader_window)
        self.btn_reader_mode.pack(padx=5, pady=5) # Ставим в удобный свободный слот

        self.btn_play = ctk.CTkButton(self.buttons_frame, text="►  Озвучить", command=self.start_speech, 
                                      fg_color="#1f8737", hover_color="#186A2B", font=("Arial", 13, "bold"))
        self.btn_play.pack(side="left", expand=True, padx=5, ipady=5)
        
        self.btn_pause = ctk.CTkButton(self.buttons_frame, text="⏸  Пауза", command=self.toggle_pause, 
                                       fg_color="#d97706", hover_color="#b45309", font=("Arial", 13, "bold"))
        self.btn_pause.pack(side="left", expand=True, padx=5, ipady=5)
        
        self.btn_stop = ctk.CTkButton(self.buttons_frame, text="⬛  Стоп", command=self.stop_speech, 
                                      fg_color="#c42b2b", hover_color="#9E2222", font=("Arial", 13, "bold"))
        self.btn_stop.pack(side="left", expand=True, padx=5, ipady=5)
        
        self.btn_export = ctk.CTkButton(self.buttons_frame, text="💾 🎵 Сохранить Аудио", command=self.export_mp3,
                                        font=("Arial", 13))
        self.btn_export.pack(side="left", expand=True, padx=5, ipady=5)

        import tkinter as tk
        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(label="✂ Вырезать", command=lambda: self.text_area.focus_get().event_generate("<<Cut>>"))
        self.context_menu.add_command(label="📋 Копировать", command=lambda: self.text_area.focus_get().event_generate("<<Copy>>"))
        self.context_menu.add_command(label="📥 Вставить", command=lambda: self.text_area.focus_get().event_generate("<<Paste>>"))
        self.text_area.bind("<Button-3>", self._show_context_menu)
        
    def refresh_interface_texts(self):
        """Мгновенно обновляет все надписи в GUI под выбранный язык self.lang"""
        self.title(self.tr("title"))
        
        # Основные кнопки панели файлов
        self.btn_load.configure(text=self.tr("btn_load"))
        self.btn_settings.configure(text=self.tr("btn_settings"))
        self.btn_add_bookmark.configure(text=self.tr("btn_bookmark"))
        self.btn_view_bookmarks.configure(text=self.tr("btn_view_bookmarks"))
        
        # Исправление: динамический перевод статуса файла книги
        if not self.last_file_path:
            self.lbl_file.configure(text=self.tr("lbl_file_empty"))
            if self.text_area.get("1.0", "end").strip() in ["Вставьте текст или откройте файл книги...", "Paste text or open a book file..."]:
                self.text_area.delete("1.0", "end")
                self.text_area.insert("1.0", self.tr("text_area_init"))
        
        # Исправление: динамический перевод метки Прогресса/Паузы
        current_progress_text = self.progress_label.cget("text")
        if "Прогресс:" in current_progress_text or "Progress:" in current_progress_text:
            pct = current_progress_text.split()[-1]
            self.progress_label.configure(text=f"{'Progress:' if self.lang == 'EN' else 'Прогресс:'} {pct}")
        elif "Пауза:" in current_progress_text or "Paused:" in current_progress_text:
            tail = current_progress_text.split(":")[-1]
            self.progress_label.configure(text=f"{self.tr('progress_pause')}{tail.strip()}")
        elif "Чтение чанков:" in current_progress_text or "Reading chunks:" in current_progress_text:
            tail = current_progress_text.split(":")[-1]
            self.progress_label.configure(text=f"{self.tr('progress_read')}{tail.strip()}")
            
        # Панель темпа речи
        try:
            val_to_check = int(str(self.speed).replace("%", "").replace("+", ""))
        except:
            val_to_check = 0
        sign = "+" if val_to_check > 0 else ""
        self.speed_label.configure(text=f"{self.tr('speed_label')}{sign}{val_to_check}")
        
        # Медиа-кнопки нижнего ряда
        self.btn_reader_mode.configure(text=self.tr("btn_reader_mode"))
        self.btn_play.configure(text=self.tr("btn_play"))
        self.btn_stop.configure(text=self.tr("btn_stop"))
        self.btn_export.configure(text=self.tr("btn_export"))
        
        # Перевод текста на кнопке паузы с учетом её состояния
        if self.is_paused:
            self.btn_pause.configure(text=self.tr("btn_resume"))
        else:
            self.btn_pause.configure(text=self.tr("btn_pause"))
            
        # Контекстное меню (Вырезать / Копировать / Вставить)
        if hasattr(self, 'context_menu'):
            try:
                self.context_menu.entryconfigure(0, label=self.tr("menu_cut"))
                self.context_menu.entryconfigure(1, label=self.tr("menu_copy"))
                self.context_menu.entryconfigure(2, label=self.tr("menu_paste"))
            except: pass
            
        # Исправление: динамический перевод заголовка окна чтения (суфлера)
        if hasattr(self, 'reader_win') and self.reader_win.winfo_exists():
            self.reader_win.title(self.tr("reader_title"))
            self.reader_win.btn_add_bookmark.configure(text=self.tr("reader_bookmark"))
            if self.reader_text.get("1.0", "end").strip() in ["Вставьте текст или откройте файл книги...", "Paste text or open a book file..."]:
                self.reader_text.delete("1.0", "end")
                self.reader_text.insert("1.0", self.tr("text_area_init"))

        if hasattr(self, 'lbl_demagog_title'):
            self.lbl_demagog_title.configure(text=self.tr("lbl_demagog"))

    def _show_context_menu(self, event):
        self.context_menu.tk_popup(event.x_root, event.y_root)

    def add_bookmark(self):
        """
        v1.8.5: Регистрирует текущую позицию чтения в базе esr_bookmarks.json.
        Поддерживает вызовы из главного окна и дочернего экрана суфлёра (ESRReaderWindow).
        Использует глобальный индекс чанка в ОЗУ для защиты от дублирования позиций.
        """
        if not self.virtual_chunks:
            messagebox.showinfo("Инфо", "Текст для создания закладки отсутствует!")
            return

        # === ИСПРАВЛЕННЫЙ ПЕРЕХВАТ ИЗ ОКНА ЧТЕНИЯ v1.8.5 ===
        if hasattr(self, '_reader_selected_text') and self._reader_selected_text:
            search_phrase = self._reader_selected_text[:70] # Берем уникальный кусок текста
            self._reader_selected_text = None # Сразу сбрасываем маркер
            
            if self.virtual_chunks:
                for idx, chunk in enumerate(self.virtual_chunks):
                    # Извлекаем текст строго по схеме: из индекса 1, если это tuple
                    chunk_text = chunk[1] if isinstance(chunk, tuple) else str(chunk)
                    
                    if search_phrase in chunk_text:
                        # Нашли точное совпадение в ОЗУ! Фиксируем правильный новый чанк
                        self.current_chunk_index = idx
                        
                        # Принудительно обновляем индикатор Маркера и оранжевую подсветку на главном экране
                        if hasattr(self, '_update_ui_progress_by_chunk'):
                            self._update_ui_progress_by_chunk()
                        elif hasattr(self, '_highlight_current_chunk'):
                            self._highlight_current_chunk()
                        break
        # ==============================================================

        try:
            # Определяем название источника (книга или первые слова копипаста)
            if self.last_file_path:
                file_name = os.path.basename(self.last_file_path)
            else:
                # Если это копипаст, берем первые 4 слова из самого первого абзаца текста
                first_para = self.paragraphs[0] if self.paragraphs else ""
                words = first_para.split()[:4]
                file_name = f"📝 Текст: {' '.join(words)}..." if words else "📝 Вставленный текст"
            
            current_chunk = self.virtual_chunks[self.current_chunk_index]
            chunk_text = current_chunk[1] if isinstance(current_chunk, tuple) else str(current_chunk)
            
            bookmarks = {}
            if os.path.exists(self.bookmarks_file):
                with open(self.bookmarks_file, "r", encoding="utf-8") as f:
                    try: bookmarks = json.load(f)
                    except: bookmarks = {}
                    
            # ... Дальше идет неизмененный оригинальный код записи в файл JSON ...
                
            new_bookmark = {
                "index": self.current_chunk_index,
                "text": chunk_text[:60] + "..." if len(chunk_text) > 60 else chunk_text
            }
            
            # === ИСПРАВЛЕНИЕ v1.8.5: Защита от KeyError, если книга открыта впервые ===
            if file_name not in bookmarks:
                bookmarks[file_name] = []
            
            if any(b["index"] == self.current_chunk_index for b in bookmarks[file_name]):
                messagebox.showinfo("Инфо", "Закладка на этом месте уже существует!")
                return
                
            bookmarks[file_name].append(new_bookmark)
            with open(self.bookmarks_file, "w", encoding="utf-8") as f:
                json.dump(bookmarks, f, ensure_ascii=False, indent=4)
                
            messagebox.showinfo(self.tr("msg_success"), f"{self.tr('msg_bm_added')}{self.current_chunk_index}")
        except Exception as e:
            messagebox.showerror("Ошибка закладок", f"Не удалось сохранить: {e}")

    def open_bookmarks(self):
        """Открывает отдельное GUI-окно со списком закладок текущей книги/копипаста."""
        if not self.virtual_chunks:
            messagebox.showinfo("Инфо", "Сначала откройте книгу или вставьте текст!")
            return
            
        # Определяем текущее название источника так же, как при сохранении
        if self.last_file_path:
            file_name = os.path.basename(self.last_file_path)
        else:
            first_para = self.paragraphs[0] if self.paragraphs else ""
            words = first_para.split()[:4]
            file_name = f"📝 Текст: {' '.join(words)}..." if words else "📝 Вставленный текст"
            
        self.current_bm_file_key = file_name # Запоминаем ключ для удаления
        
        bookmarks_list = []
        if os.path.exists(self.bookmarks_file):
            try:
                with open(self.bookmarks_file, "r", encoding="utf-8") as f:
                    all_bookmarks = json.load(f)
                    bookmarks_list = all_bookmarks.get(file_name, [])
            except Exception as e:
                print(f"[ОШИБКА ЧТЕНИЯ ЗАКЛАДОК]: {e}")
                
        if not bookmarks_list:
            messagebox.showinfo("Инфо", f"Для «{file_name}» пока нет сохраненных закладок.")
            return

        self.bm_window = ctk.CTkToplevel(self)
        self.bm_window.title(self.tr("bm_title")) # ИСПРАВЛЕНО
        self.bm_window.geometry("500x420")
        self.bm_window.resizable(False, False)
        self.bm_window.geometry(f"+{self.winfo_x() + 100}+{self.winfo_y() + 120}")
        self.bm_window.attributes("-topmost", True)

        lbl_info = ctk.CTkLabel(self.bm_window, text=self.tr("bm_info"), font=("Arial", 12, "bold")) # ИСПРАВЛЕНО
        lbl_info.pack(pady=(10, 5), padx=20, anchor="w")

        self.bm_display = ctk.CTkTextbox(self.bm_window, width=460, height=280, font=("Arial", 13), cursor="hand2")
        self.bm_display.pack(padx=20, pady=5, fill="both", expand=True)

        # Кнопка удаления внизу окна
        self.btn_del_bm = ctk.CTkButton(
            self.bm_window, 
            text=self.tr("bm_btn_del"), # ИСПРАВЛЕНО
            command=self.delete_selected_bookmark, 
            fg_color="#c42b2b", 
            hover_color="#9E2222", 
            font=("Arial", 13, "bold")
        )
        self.btn_del_bm.pack(pady=10)

        # === ИСПРАВЛЕНИЕ v1.8.5: Вывод собранных закладок на экран ===
        lines = []
        self.bm_index_map = {}

        for i, bm in enumerate(bookmarks_list):
            idx = bm.get("index", 0)
            txt = bm.get("text", "")
            # Формируем красивую локализованную строку
            line_text = f" [📖 {file_name}] | {self.tr('bm_pos')} {idx} ➔ {txt}\n"
            lines.append(line_text)
            self.bm_index_map[i + 1] = idx

        # Включаем режим редактирования виджета, чтобы залить текст
        self.bm_display.configure(state="normal")
        self.bm_display.delete("1.0", "end")
        
        if lines:
            self.bm_display.insert("1.0", "".join(lines))
        else:
            self.bm_display.insert("1.0", "Список закладок пуст.")

        # Возвращаем режим «только чтение», чтобы пользователь не мог стирать текст курсором
        self.bm_display.configure(state="disabled")

        # Переменная для хранения строки, которую выделили одиночным щелчком
        self.selected_bm_line = None

        self.bm_display.bind("<Button-1>", self._on_bookmark_single_click)
        self.bm_display.bind("<Double-Button-1>", self._on_bookmark_double_click)

    def _on_bookmark_single_click(self, event):
        """Фиксирует строку при одиночном клике и визуально подсвечивает её."""
        try:
            # Даем права на изменение стилей текста
            self.bm_display.configure(state="normal")
            
            # Настраиваем цвет фона для выделенной строки (светло-голубой)
            self.bm_display.tag_config("selected_line", background="#e0f2fe", foreground="black")
            # Сбрасываем прошлую подсветку со всего поля
            self.bm_display.tag_remove("selected_line", "1.0", "end")
            
            # Вычисляем строку клика
            click_pos = self.bm_display.index(f"@{event.x},{event.y}")
            self.selected_bm_line = int(click_pos.split('.')[0])
            
            # Накладываем подсветку на всю выбранную строку от начала до конца
            start_idx = f"{self.selected_bm_line}.0"
            end_idx = f"{self.selected_bm_line}.end"
            self.bm_display.tag_add("selected_line", start_idx, end_idx)
            
        except Exception as e:
            print(f"[ОШИБКА ВЫДЕЛЕНИЯ СТРОКИ]: {e}")
            self.selected_bm_line = None
            
        # Возвращаем режим «только чтение»
        self.bm_display.configure(state="disabled")

    def _on_bookmark_double_click(self, event):
        """Обрабатывает двойной щелчок по закладке и перемещает чтение."""
        try:
            click_pos = self.bm_display.index(f"@{event.x},{event.y}")
            line_num = int(click_pos.split('.')[0])
            target_chunk_idx = self.bm_index_map.get(line_num)
            
            if target_chunk_idx is not None and target_chunk_idx < len(self.virtual_chunks):
                self.current_chunk_index = target_chunk_idx
                para_idx, chunk_text = self.virtual_chunks[self.current_chunk_index]
                
                total_chunks = len(self.virtual_chunks)
                prog = self.current_chunk_index / total_chunks if total_chunks else 0
                self._update_ui_progress(prog, f"Закладка: {int(prog*100)}% ({self.current_chunk_index}/{total_chunks})")
                
                if para_idx < len(self.paragraphs):
                    paragraph_text = self.paragraphs[para_idx]
                    if paragraph_text:
                        self.text_area.tag_remove("para_highlight", "1.0", "end")
                        idx = self.text_area._textbox.search(paragraph_text, "1.0", stopindex="end")
                        if idx:
                            end_idx = f"{idx} + {len(paragraph_text)} chars"
                            self.text_area.tag_add("para_highlight", idx, end_idx)
                            self.text_area._textbox.see(idx)
                
                if self.is_speaking and not self.is_paused:
                    self.is_speaking = False
                    self.is_prefetching = False
                    pygame.mixer.music.stop()
                    pygame.mixer.stop()
                    self.after(100, self.start_speech)
                else:
                    self.save_state()

            # --- СИНХРОНИЗАЦИЯ ОКНА ЧТЕНИЯ v1.8.5 ---
            if hasattr(self, 'reader_win') and self.reader_win.winfo_exists():
                self.reader_win.sync_with_active_chunk()

                self.bm_window.destroy()
        except Exception as e:
            print(f"[ОШИБКА ПЕРЕХОДА ПО ЗАКЛАДКЕ]: {e}")

    def delete_selected_bookmark(self):
        """Физически удаляет выбранную закладку из файла esr_bookmarks.json и обновляет окно."""
        if not hasattr(self, 'selected_bm_line') or self.selected_bm_line is None:
            messagebox.showinfo("Инфо", "Сначала выберите закладку одиночным кликом мыши!")
            return
            
        try:
            # Получаем индекс чанка, который нужно удалить, по номеру выделенной строки
            target_chunk_idx = self.bm_index_map.get(self.selected_bm_line)
            if target_chunk_idx is None:
                return
                
            if os.path.exists(self.bookmarks_file):
                with open(self.bookmarks_file, "r", encoding="utf-8") as f:
                    all_bookmarks = json.load(f)
                    
                # Ищем список закладок для текущей книги/копипаста
                file_key = self.current_bm_file_key
                if file_key in all_bookmarks:
                    # Фильтруем список, исключая удаляемую позицию
                    old_list = all_bookmarks[file_key]
                    new_list = [b for b in old_list if b.get("index") != target_chunk_idx]
                    
                    if len(old_list) == len(new_list):
                        return # Ничего не изменилось
                        
                    all_bookmarks[file_key] = new_list
                    
                    # Если закладок для этой книги больше не осталось, удаляем сам ключ книги
                    if not new_list:
                        del all_bookmarks[file_key]
                        
                    # Записываем обновленные данные на диск
                    with open(self.bookmarks_file, "w", encoding="utf-8") as f:
                        json.dump(all_bookmarks, f, ensure_ascii=False, indent=4)

                    # === ИСПРАВЛЕНИЕ v1.8.5: Убираем неопределенный флаг success ===
                    messagebox.showinfo(self.tr("msg_success"), self.tr("msg_bm_deleted"))

                    # Закрываем окно закладок, чтобы обновить данные
                    self.bm_window.destroy()

                    # Если в этой книге еще остались другие закладки, сразу переоткрываем окно с новым списком
                    if new_list:
                        self.after(100, self.open_bookmarks)
        except Exception as e:
            messagebox.showerror("Ошибка удаления", f"Не удалось удалить: {e}")

    def open_settings(self):
        if self.settings_window is not None and self.settings_window.winfo_exists():
            self.settings_window.focus()
            return
            
        self.settings_window = ctk.CTkToplevel(self)
        self.settings_window.title("⚙ Настройки ESR")
        self.settings_window.geometry("440x380")
        self.settings_window.resizable(False, False)
        self.settings_window.geometry(f"+{self.winfo_x() + 140}+{self.winfo_y() + 130}")
        self.settings_window.attributes("-topmost", True)
        
        if os.path.exists("icon.ico"):
            try:
                self.settings_window.wm_iconbitmap("icon.ico")
            except: pass
            
        self._build_settings_tabs()

    # 📦 БЛОК 1.2: Логика и наполнение вкладок настроек
    # --- Что внутри: Изолированные методы создания вкладок, тексты встроенной справки и правовой лицензии MIT.
    # --- ПАМЯТЬ ДЛЯ ГРАФИЧЕСКИХ НАСТРОЕК ---
    # Что изменилось: Списки тем и фонов теперь жестко инициализируются теми значениями, которые лежат в self.theme и self.text_bg. При выборе они обновляют эти переменные.(05.07.2026 01:10)
    def _build_settings_tabs(self):
        # Создаем вкладки на базе локальной переменной tabview
        tabview = ctk.CTkTabview(self.settings_window, width=410, height=350)
        tabview.pack(padx=15, pady=10, fill="both", expand=True)
        
        tab_interface_name = self.tr("tab_interface")
        tab_help_name = self.tr("tab_help")
        tab_about_name = self.tr("tab_about")
        
        tabview.add(tab_interface_name)
        tabview.add(tab_help_name)
        tabview.add(tab_about_name)
        
        # --- ВКЛАДКА: ИНТЕРФЕЙС (ОПТИМИЗАЦИЯ СЕТКИ v1.8.5) ---
        # 1. Блок выбора языка
        lang_frame = ctk.CTkFrame(tabview.tab(tab_interface_name))
        lang_frame.pack(fill="x", padx=10, pady=2) # Уменьшен pady
        
        lang_label = ctk.CTkLabel(lang_frame, text=self.tr("lbl_lang"), font=("Arial", 12, "bold"))
        lang_label.pack(anchor="w", padx=15, pady=2)
        
        lang_combo = ctk.CTkComboBox(lang_frame, values=["Русский (RU)", "English (EN)"], command=self.change_interface_language, width=220, height=26)
        lang_combo.set("Русский (RU)" if self.lang == "RU" else "English (EN)")
        lang_combo.pack(padx=15, pady=2, anchor="w")
        
        # 2. Блок выбора темы
        theme_frame = ctk.CTkFrame(tabview.tab(tab_interface_name))
        theme_frame.pack(fill="x", padx=10, pady=2) # Уменьшен pady
        
        theme_label = ctk.CTkLabel(theme_frame, text=self.tr("lbl_theme"), font=("Arial", 12, "bold"))
        theme_label.pack(anchor="w", padx=15, pady=2)
        
        theme_combo = ctk.CTkComboBox(theme_frame, values=["Светлая (Light)", "Тёмная (Dark)"], command=self.change_theme, width=220, height=26)
        theme_combo.set(self.theme)
        theme_combo.pack(padx=15, pady=2, anchor="w")
        
        # 3. Блок выбора фона текста (Локализованный в v1.8.5)
        bg_frame = ctk.CTkFrame(tabview.tab(tab_interface_name))
        bg_frame.pack(fill="x", padx=10, pady=2)

        bg_label = ctk.CTkLabel(bg_frame, text=self.tr("lbl_bg"), font=("Arial", 12, "bold"))
        bg_label.pack(anchor="w", padx=15, pady=2)

        # Массив значений формируется динамически на основе текущего языка
        bg_values = [self.tr("bg_standard"), self.tr("bg_sepia"), self.tr("bg_oled")]

        self.bg_combo = ctk.CTkComboBox(bg_frame, values=bg_values, command=self.change_text_bg, width=220, height=26)

        # Устанавливаем текущее выбранное значение
        if self.text_bg in ["«Стандартный»", "Default Style"]:
            self.bg_combo.set(self.tr("bg_standard"))
        elif self.text_bg in ["«Книжная Сепия»", "Book Sepia"]:
            self.bg_combo.set(self.tr("bg_sepia"))
        elif self.text_bg in ["«Ночной OLED»", "Night OLED"]:
            self.bg_combo.set(self.tr("bg_oled"))
        else:
            self.bg_combo.set(self.tr("bg_standard"))

        self.bg_combo.pack(padx=15, pady=2, anchor="w")
        
        # 4. Блок пользовательского словаря
        dict_frame = ctk.CTkFrame(tabview.tab(tab_interface_name))
        dict_frame.pack(fill="x", padx=10, pady=2) # Уменьшен pady
        
        dict_label = ctk.CTkLabel(dict_frame, text="Управление личными ударениями:" if self.lang == "RU" else "Personal Pronunciation Dict:", font=("Arial", 12, "bold"))
        dict_label.pack(anchor="w", padx=15, pady=2)
        
        btn_open_dict = ctk.CTkButton(dict_frame, text=self.tr("btn_open_dict"), command=self.open_user_dictionary, width=220, height=26, fg_color="#10b981", hover_color="#059669")
        btn_open_dict.pack(padx=15, pady=4, anchor="w")
        
        # --- ВКЛАДКА: ПОМОЩЬ (v1.8.5 — Интерактивные ссылки на README) ---
        help_box = ctk.CTkTextbox(tabview.tab(tab_help_name), width=380, height=260, font=("Arial", 12))
        help_box.pack(padx=5, pady=5, fill="both", expand=True)

        # Настраиваем стили и привязки для ссылок через базовое ядро Tkinter
        help_box._textbox.tag_config("link", foreground="#115e59", underline=True)
        help_box._textbox.tag_bind("link", "<Enter>", lambda e: help_box.configure(cursor="hand2"))
        help_box._textbox.tag_bind("link", "<Leave>", lambda e: help_box.configure(cursor=""))
        
        # Лямбда-обработчики для безопасного открытия файлов по клику
        help_box._textbox.tag_bind("readme_ru", "<Button-1>", lambda e: self._open_local_readme("README.md"))
        help_box._textbox.tag_bind("readme_en", "<Button-1>", lambda e: self._open_local_readme("README_EN.md"))
        
        if self.lang == "RU":
            help_box.insert("1.0", (
                "ОСНОВНЫЕ ФУНКЦИИ И ВОЗМОЖНОСТИ:\n"
                "• Гибридный движок озвучки: Поддержка реалистичных онлайн-голосов Microsoft Edge-TTS "
                "и встроенных системных офлайн-синтезаторов SAPI5.\n"
                "• RAM-кэширование (Без диска): Аудиопоток генерируется напрямую в оперативную память "
                "через BytesIO-буферы. Программа больше не создаёт временные файлы на жёстком диске, "
                "что исключает износ SSD и минимизирует задержки.\n"
                "• Асинхронный префетчер: Умный фоновый предзаказ чанков (глубина: 4 абзаца вперёд). "
                "Обеспечивает бесшовное воспроизведение без пауз на загрузку текста.\n"
                "• Синхронное окно чтения: Дополнительный экран-суфлёр, который полностью дублирует "
                "позицию текста главного окна и мгновенно адаптируется под выбранную световую схему.\n"
                "• Умная система закладок: Локальное сохранение позиций в JSON-базу по уникальным индексам "
                "чангов, защищённое от дубликатов и сбоев при первом открытии книг.\n"
                "• Цветовые пресеты: Быстрое переключение интерфейса между Стандартным, Книжной Сепией "
                "и Ночным OLED режимами для снижения нагрузки на зрение.\n\n"
                "ГОРЯЧИЕ КЛАВИШИ УПРАВЛЕНИЯ (СУФЛЁР):\n"
                "• Прокрутка и навигация полностью синхронизированы с основным окном.\n"
                "• Кнопка «+ Закладки» мгновенно фиксирует чанк прямо во время прослушивания.\n\n"
                "Система полностью автономна, оптимизирована для Windows 11 x64 и готова к работе."
                "ДОКУМЕНТАЦИЯ ПРОЕКТА:\n"
                "👉 Открыть полную документацию: "
            ))
            # Вставляем кликабельный текст с тегом ссылки
            help_box._textbox.insert("end", "README.md", ("link", "readme_ru"))
        else:
            help_box.insert("1.0", (
                "KEY FEATURES & CAPABILITIES:\n"
                "• Hybrid TTS Engine: Supports realistic online voices via Microsoft Edge-TTS "
                "as well as built-in native offline SAPI5 system synthesizers.\n"
                "• RAM-Caching (Diskless): Audio streams are generated directly into RAM "
                "using BytesIO buffers. The application no longer creates temporary files on the disk, "
                "eliminating SSD wear and reducing playback latency.\n"
                "• Asynchronous Prefetcher: Smart background chunk pre-ordering (buffer depth: 4 paragraphs ahead). "
                "Ensures seamless playback without delays during text loading.\n"
                "• Synchronous Reader Window: An auxiliary prompter screen that fully mirrors "
                "the text position of the main window and instantly adapts to the selected color scheme.\n"
                "• Smart Bookmark System: Locally saves reading positions into a JSON database using unique "
                "chunk indices, fully protected against duplicates and bugs when opening books for the first time.\n"
                "• Visual Color Presets: Quick interface toggling between Standard, Book Sepia, "
                "and Night OLED modes to significantly reduce eye strain.\n\n"
                "CONTROL HOTKEYS (PROMPTER):\n"
                "• Scrolling and navigation are fully synchronized with the main interface window.\n"
                "• The '+ Bookmark' button instantly pins the current chunk during active listening.\n\n"
                "The system is completely portable, fully optimized for Windows 11 x64, and ready to go."
                "PROJECT DOCUMENTATION:\n"
                "👉 Open full documentation: "
            ))
            # Вставляем английский кликабельный текст
            help_box._textbox.insert("end", "README_EN.md", ("link", "readme_en"))

        # help_box = ctk.CTkTextbox(tabview.tab(tab_help_name), width=380, height=260, font=("Arial", 12))
        # help_box.pack(padx=5, pady=5, fill="both", expand=True)
        # help_box.insert("1.0", help_text)
        help_box.configure(state="disabled")
        
        # --- ВКЛАДКА: О ПРОГРАММЕ ---
        if self.lang == "RU":
            about_text = (
                "Portable EdgeStream Reader (ESR) v1.8.5-beta\n"
                "Разработчик софта: timurer & Гай(AI Google) (2026)\n\n"
                "----------------------------------------------------------------------\n"
                "ЛИЦЕНЗИОННОЕ СОГЛАШЕНИЕ (MIT License):\n\n"
                "Программное обеспечение предоставляется БЕСПЛАТНО, «КАК ЕСТЬ» (As Is),\n"
                "без каких-либо гарантий, явно выраженных или подразумеваемых.\n\n"
                "Разработчики не несут никакой ответственности за любые претензии, "
                "убытки или иные обязательства, возникшие в результате использования "
                "данного программного обеспечения."
            )
        else:
            about_text = (
                "Portable EdgeStream Reader (ESR) v1.8.5-beta\n"
                "Software Developer: timurer & Guy(AI Google) (2026)\n\n"
                "----------------------------------------------------------------------\n"
                "LICENSE AGREEMENT (MIT License):\n\n"
                "The software is provided FREE of charge, 'AS IS', without warranty\n"
                "of any kind, express or implied.\n\n"
                "The developers shall not be liable for any claims, damages or other\n"
                "liability arising from the use of this software."
            )
        about_box = ctk.CTkTextbox(tabview.tab(tab_about_name), width=380, height=260, font=("Arial", 11))
        about_box.pack(padx=5, pady=5, fill="both", expand=True)
        about_box.insert("1.0", about_text)
        about_box.configure(state="disabled")

    def change_interface_language(self, choice):
        """Обработчик выбора языка в комбобоксе: перерисовывает все окна приложения"""
        self.lang = "RU" if "Русский" in choice else "EN"
        
        # 1. Обновляем все тексты главного окна и окна чтения
        self.refresh_interface_texts()
        
        # 2. Пересохраняем новый конфиг на диск
        self.save_state()
        
        # 3. Полностью уничтожаем и перерисовываем содержимое окна настроек на лету
        if self.settings_window is not None and self.settings_window.winfo_exists():
            for widget in self.settings_window.winfo_children():
                widget.destroy()
            self._build_settings_tabs()

    def change_theme(self, choice):
        self.theme = choice
        # Пишем безопасное условие, работающее независимо от выбранного языка интерфейса
        if "Тёмная" in choice or "Dark" in choice:
            ctk.set_appearance_mode("Dark")
            self.text_area.tag_config("para_highlight", background="#115e59", foreground="white")
        else:
            ctk.set_appearance_mode("Light")
            self.text_area.tag_config("para_highlight", background="#ccfbf1", foreground="black")
        
        # Обновляем фон текстового поля, считывая значение из комбобокса напрямую
        if hasattr(self, 'bg_combo') and self.bg_combo is not None:
            self.change_text_bg(self.bg_combo.get())
        else:
            self.change_text_bg(self.text_bg)
        
        self.save_state()

    def change_text_bg(self, choice):
        """
        Переключает цветовую палитру интерфейса (Стандартный, Сепия, OLED).
        v1.8.5: Динамически транслирует выбранную схему (HEX-коды) в открытое 
        окно чтения (ESRReaderWindow) для обеспечения визуальной синхронности окон.
        """
        if choice in ["«Книжная Сепия»", "Book Sepia"]:
            self.text_bg = "«Книжная Сепия»"
            fg, txt, highlight_bg, highlight_fg = "#f4ecd8", "#5b4636", "#e4d1b3", "#5b4636"
        elif choice in ["«Ночной OLED»", "Night OLED"]:
            self.text_bg = "«Ночной OLED»"
            fg, txt, highlight_bg, highlight_fg = "#000000", "#d1d5db", "#1f2937", "#ffffff"
        else:
            self.text_bg = "«Стандартный»"
            if ctk.get_appearance_mode() == "Dark":
                fg, txt, highlight_bg, highlight_fg = "#1f2937", "white", "#115e59", "white"
            else:
                fg, txt, highlight_bg, highlight_fg = "white", "black", "#ccfbf1", "black"

        # Применяем цвета к главному окну
        self.text_area.configure(fg_color=fg, text_color=txt)
        self.text_area.tag_config("para_highlight", background=highlight_bg, foreground=highlight_fg)

        # === ДОРАБОТКА v1.8.5: Перекраска дочернего окна чтения на лету ===
        if hasattr(self, 'reader_win') and self.reader_win is not None and self.reader_win.winfo_exists():
            self.reader_win.reader_text.configure(fg_color=fg, text_color=txt)
            self.reader_win.configure(fg_color=ctk.ThemeManager.theme["CTkFrame"]["fg_color"] if ctk.get_appearance_mode() == "Light" else ctk.ThemeManager.theme["CTkFrame"]["top_highlight"])

        self.save_state()

    def change_voice(self, choice):
        self.voice = self.voice_map[choice]
        self.save_state()

    def on_speed_change(self, value):
        """Безопасный обработчик изменения ползунка скорости (Шкала от -10 до +10)"""
        try:
            current_val = int(round(float(value)))
            sign = "+" if current_val > 0 else ""
            self.speed = current_val
            
            if hasattr(self, 'speed_label') and self.speed_label is not None:
                self.speed_label.configure(text=f"Скорость темпа: {sign}{current_val}")
        except Exception as e:
            print(f"[ОШИБКА ОБНОВЛЕНИЯ ПОЛЗУНКА]: {e}")
        
    #📦 1.3: Модуль пользовательского словаря ударений
    #Что внутри: Полная интерактивная логика Toplevel-окна словаря, контекстные меню по правому клику мыши для ввода спецсимволов, пошаговое чтение/запись JSON-файла.
    # --- БЛОК 1.3: ИНТЕРФЕЙС И ЛОГИКА ПОЛЬЗОВАТЕЛЬСКОГО СЛОВАРЯ ---
    def open_user_dictionary(self):
        # --- НОВАЯ СИНХРОНИЗАЦИЯ ОКН v1.8.5 ---
        # Если окно настроек открыто, автоматически закрываем его перед открытием словаря
        if self.settings_window is not None and self.settings_window.winfo_exists():
            self.settings_window.destroy()
            self.settings_window = None # Очищаем ссылку для безопасности

        # --- ЗАЩИТА v1.8.5: Гарантируем наличие переменной словаря ---
        if not hasattr(self, 'user_dict'):
            self.user_dict = {}
            
        # Проверяем, не открыто ли окно словаря уже
        if hasattr(self, 'dict_win') and self.dict_win.winfo_exists():
            self.dict_win.focus()
            return
            
        self.dict_win = ctk.CTkToplevel(self)
        self.dict_win.title(self.tr("dict_win_title")) 
        self.dict_win.geometry("450x480")
        
        # --- УСТАНОВКА РОДНОЙ ИКОНКИ ДЛЯ ОКНА СЛОВАРЯ v1.8.5 ---
        if os.path.exists("icon.ico"):
            try:
                self.dict_win.wm_iconbitmap("icon.ico")
            except: 
                pass
        else:
            # Если файла иконки физически нет на диске, убираем стандартное перо Tkinter
            self.dict_win.after(200, lambda: self.dict_win.iconbitmap(default=""))
            
        self.dict_win.grab_set()
        
        # Текстовое поле вывода правил
        self.dict_textbox = ctk.CTkTextbox(self.dict_win, width=410, height=200, font=("Arial", 12))
        self.dict_textbox.pack(padx=20, pady=15, fill="both", expand=True)
        
        # Функция обновления списка внутри текстового окна
        def refresh_dict_display():
            self.dict_textbox.configure(state="normal")
            self.dict_textbox.delete("1.0", "end")
            # ИСПРАВЛЕНО: возвращаем имя self.user_dict из оригинального кода
            if not self.user_dict:
                self.dict_textbox.insert("1.0", self.tr("dict_empty"))
            else:
                for idx, (k, v) in enumerate(self.user_dict.items(), 1):
                    self.dict_textbox.insert("end", f"{idx}. {k} ➔ {v}\n")
            self.dict_textbox.configure(state="disabled")
            
        refresh_dict_display()
        
        # Поле ввода нового правила
        entry_frame = ctk.CTkFrame(self.dict_win, fg_color="transparent")
        entry_frame.pack(fill="x", padx=20, pady=5)
        
        dict_entry = ctk.CTkEntry(entry_frame, placeholder_text=self.tr("dict_placeholder"), width=260)
        dict_entry.pack(side="left", padx=(0, 10), fill="x", expand=True)
        
        # Функция добавления правила
        def add_rule_click():
            val = dict_entry.get().strip()
            if "=" in val:
                k, v = val.split("=", 1)
                # ИСПРАВЛЕНО: пишем в оригинальный словарь self.user_dict
                self.user_dict[k.strip()] = v.strip()
                self.save_user_dictionary()
                refresh_dict_display()
                dict_entry.delete(0, "end")
                
        btn_add = ctk.CTkButton(entry_frame, text=self.tr("dict_btn_add"), command=add_rule_click, width=120)
        btn_add.pack(side="right")
        
        # Блок удаления правил
        del_frame = ctk.CTkFrame(self.dict_win, fg_color="transparent")
        del_frame.pack(fill="x", padx=20, pady=15)
        
        # Метка выбранного слова
        self.lbl_selected_word = ctk.CTkLabel(
            del_frame, 
            text=f"{self.tr('dict_selected')}{self.tr('dict_none')}", 
            font=("Arial", 12)
        )
        self.lbl_selected_word.pack(side="left", anchor="w")
        
        # Поле для ввода номера/слова на удаление
        del_entry = ctk.CTkEntry(del_frame, placeholder_text="№", width=50)
        del_entry.pack(side="left", padx=10)
        
        # Функция удаления правила
        def del_rule_click():
            num_str = del_entry.get().strip()
            if num_str.isdigit():
                idx = int(num_str) - 1
                # ИСПРАВЛЕНО: читаем из оригинального словаря self.user_dict
                keys = list(self.user_dict.keys())
                if 0 <= idx < len(keys):
                    del self.user_dict[keys[idx]]
                    self.save_user_dictionary()
                    refresh_dict_display()
                    del_entry.delete(0, "end")
                    self.lbl_selected_word.configure(text=f"{self.tr('dict_selected')}{self.tr('dict_none')}")
                    
        btn_del = ctk.CTkButton(
            del_frame, 
            text=self.tr("dict_btn_del"), 
            command=del_rule_click, 
            width=120, 
            fg_color="#dc2626", 
            hover_color="#b91c1c"
        )
        btn_del.pack(side="right")

    def _on_dict_line_click(self, event):
        # Алгоритм определяет номер строки по координате пикселя клика мыши
        self.dict_display.configure(state="normal")
        try:
            idx = self.dict_display.index(f"@{event.x},{event.y}")
            line_num = idx.split(".")[0]
            line_text = self.dict_display.get(f"{line_num}.0", f"{line_num}.end").strip()
            
            if "➔" in line_text:
                word = line_text.split("➔")[0].strip()
                self.selected_word_to_delete = word
                self.lbl_selected_word.configure(text=f"Выбрано слово: {word}", text_color="#10b981")
        except: pass
        self.dict_display.configure(state="disabled")

    def _get_user_dict(self):
        if not os.path.exists(self.dict_file): return {}
        try:
            with open(self.dict_file, "r", encoding="utf-8") as f: return json.load(f)
        except: return {}

    def _refresh_dict_display(self):
        self.dict_display.configure(state="normal")
        self.dict_display.delete("1.0", "end")
        self.selected_word_to_delete = ""
        self.lbl_selected_word.configure(text="Выбрано слово: [нет]", text_color=("black", "white"))
        d = self._get_user_dict()
        if not d:
            self.dict_display.insert("1.0", "Словарь пуст. Добавьте свои правила в формате слово=замена.")
        else:
            lines = [f"{k} ➔ {v}" for k, v in d.items()]
            self.dict_display.insert("1.0", "\n".join(lines))
        self.dict_display.configure(state="disabled")

    def add_dict_item(self):
        raw = self.dict_input.get().strip()
        if "=" not in raw:
            messagebox.showerror("Ошибка", "Неверный формат! Используйте знак равенства (=).")
            return
        k, v = raw.split("=", 1)
        k, v = k.strip().lower(), v.strip()
        if not k or not v: return
        
        d = self._get_user_dict()
        d[k] = v
        try:
            with open(self.dict_file, "w", encoding="utf-8") as f:
                json.dump(d, f, ensure_ascii=False, indent=4)
            self.dict_input.delete(0, "end")
            self._refresh_dict_display()
        except: pass

    def delete_dict_item(self):
        k = self.selected_word_to_delete
        if not k:
            messagebox.showinfo("Инфо", "Сначала выберите строку с правилом кликом мыши!")
            return
        d = self._get_user_dict()
        if k in d:
            del d[k]
            try:
                with open(self.dict_file, "w", encoding="utf-8") as f:
                    json.dump(d, f, ensure_ascii=False, indent=4)
                self._refresh_dict_display()
            except: pass

    # 📦 Блок 2: Парсинг путей, Книжные кодировки и Оптимизация пауз
    # Что внутри: Инструменты _parse_path_safely, _read_file_content, загрузчик load_file и регулярки optimize_text_pauses.
    #04.07.2026. Что изменилось: добавлен метод _clean_text_notepad_killer (наш умный чистильщик), который автоматически вызывается внутри load_file.
    #12.07.2026. Всеядный импорт книг.Новые форматы .epub и .docx
    #12.07.2026. Умные словари: Словарь Е/Ё; Парсер сокращений (Скрипт смотрит, что стоит перед точкой: если цифра (например, 1945 г.), он заменит это на года. Если перед «г.» стоит пробел и большая буква (например, г. Москва), он аккуратно расшифрует это как город. Также добавим килограммы, метры, километры и рубли.)
    def _split_text_into_smart_chunks(self, raw_text):
        raw_text = str(raw_text)
        if not raw_text.strip():
            return [], []
        
        # --- МГНОВЕННАЯ НАРЕЗКА КНИГИ БЕЗ ЗАВИСАНИЙ UI-ПОТОКА ---
        source_paragraphs = [p.strip() for p in raw_text.split("\n") if p.strip()]
        virtual_chunks = []
        
        for para_idx, para_text in enumerate(source_paragraphs):
            sentences = re.split(r'(?<=[.!?])\s+', para_text)
            temp_chunk = []
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence: continue
                temp_chunk.append(sentence)
                
                words_count = len(re.findall(r'\b\w+\b', " ".join(temp_chunk)))
                if words_count >= 4 or sentence == sentences[-1]:
                    virtual_chunks.append((para_idx, " ".join(temp_chunk)))
                    temp_chunk = []
            if temp_chunk:
                virtual_chunks.append((para_idx, " ".join(temp_chunk)))
                
        return source_paragraphs, virtual_chunks

    def _on_text_click(self, event):
        """Вычисляет чанк под курсором и переключает на него индекс чтения."""
        if not self.virtual_chunks:
            return
            
        try:
            # Получаем индекс символа (строка.символ) Tkinter под курсором мыши
            click_index = self.text_area._textbox.index(f"@{event.x},{event.y}")
            line_num, char_num = map(int, click_index.split('.'))
            
            # Получаем текст всей строки, где был совершен клик
            line_text = self.text_area._textbox.get(f"{line_num}.0", f"{line_num}.end")
            
            # Находим, какое именно предложение внутри строки было выбрано
            matched_chunk_idx = None
            for idx, (para_idx, chunk_text) in enumerate(self.virtual_chunks):
                if chunk_text in line_text:
                    start_pos = line_text.find(chunk_text)
                    end_pos = start_pos + len(chunk_text)
                    # Если координата клика падает внутрь границ чанка
                    if start_pos <= char_num <= end_pos:
                        matched_chunk_idx = idx
                        break
            
            if matched_chunk_idx is not None:
                self.current_chunk_index = matched_chunk_idx
                
                # Синхронное обновление UI прогресса
                total_chunks = len(self.virtual_chunks)
                prog = self.current_chunk_index / total_chunks if total_chunks else 0
                self._update_ui_progress(prog, f"Маркер: {int(prog*100)}% ({self.current_chunk_index}/{total_chunks})")
                
                # Мгновенно подсвечиваем родительский абзац на экране
                actual_para_idx = self.virtual_chunks[self.current_chunk_index][0]
                self._highlight_paragraph_on_screen(actual_para_idx)
                
                # Если плеер работал, мягко перезапускаем его с новой позиции
                if self.is_speaking and not self.is_paused:
                    self.is_speaking = False
                    self.is_prefetching = False
                    pygame.mixer.music.stop()
                    pygame.mixer.stop()
                    self.after(100, self.start_speech)
                else:
                    self.save_state()
        except Exception as e:
            print(f"[МАРКЕР ОШИБКА]: {e}")

    def optimize_text_pauses(self, text):
        if not text: return ""
        # Заменяем только отдельно стоящие длинные и короткие тире с пробелами
        # Дефисы внутри слов вида \w-\w (как-то, где-то) больше не заменяются на запятые!
        text = re.sub(r'\s+[-—–]\s+', ', ', text)
        
        # Остальные знаки препинания обрабатываем аккуратно
        for char in [r'\.\.\.\s*', r'\.\s*', r';\s*', r':\s*']:
            text = re.sub(char, ', ', text)
        return text

    # --- ДОРАБОТКА: ПОЛНАЯ НОРМАЛИЗАЦИЯ И ПОЛЬЗОВАТЕЛЬСКИЙ СЛОВАРЬ ---
    def _normalize_text_dictionaries(self, text):
        if not text: return ""
        
        # --- ПОДКЛЮЧЕНИЕ ЛИЧНОГО СЛОВАРЯ УДАРЕНИЙ И ЗАМЕН ---
        dict_file = "esr_user_dict.json"
        if os.path.exists(dict_file):
            try:
                with open(dict_file, "r", encoding="utf-8") as f:
                    user_dict = json.load(f)
                # Проходим по всем правилам из файла
                for search_word, replace_word in user_dict.items():
                    if search_word and replace_word:
                        # Используем регулярку, чтобы менять слово целиком вне зависимости от регистра (Регистронезависимо)
                        pattern = re.compile(r'\b' + re.escape(search_word) + r'\b', re.IGNORECASE)
                        
                        # Вспомогательная функция, чтобы сохранять заглавную букву, если слово стоит в начале предложения
                        def dict_replacer(match, rep=replace_word):
                            orig = match.group(0)
                            if orig and orig[0].isupper():
                                return rep[0].upper() + rep[1:]
                            return rep
                            
                        text = pattern.sub(dict_replacer, text)
            except Exception as e:
                print(f"[ERROR] Ошибка применения пользовательского словаря: {e}")

        # 1. Замена буквы Е на Ё
        ee_replacements = {
            r'\bвсе\b': 'всё',
            r'\bеще\b': 'ещё',
            r'\bЕще\b': 'Ещё',
            r'\bмое\b': 'моё',
            r'\bтвое\b': 'твоё',
            r'\bсвое\b': 'своё',
            r'\bнее\b': 'неё',
            r'\bчерт\b': 'чёрт',
        }
        for pattern, repl in ee_replacements.items():
            text = re.sub(pattern, repl, text)

        # Математическая формула склонения числительных
        def get_declension(number_str, forms):
            try:
                num = int(number_str)
                if 11 <= num % 100 <= 14:
                    return forms
                remainder = num % 10
                if remainder == 1:
                    return forms
                if remainder == 2 or remainder == 3 or remainder == 4:
                    return forms
                return forms
            except:
                return forms

        # 2. УМНЫЙ ПАРСЕР СОКРАЩЕНИЙ ГОДОВ
        def repl_year_in(match):
            return f"{match.group(1)} {match.group(2)} году"
        text = re.sub(r'\b([вВ])\s+(\d+)\s*г\b\.?', repl_year_in, text)
        text = re.sub(r'\b([вВ])\s+(\d+)\s*гг\b\.?', r'\1 \2 годах', text)

        def repl_year_gen(match):
            num_str = match.group(1)
            form = get_declension(num_str, ["год", "года", "лет"])
            return f"{num_str} {form}"
        text = re.sub(r'(\d+)\s*г\b\.?', repl_year_gen, text)
        text = re.sub(r'(\d+)\s*гг\b\.?', r'\1 годов', text)
        
        text = re.sub(r'\bg\.\s*(?=[А-ЯЁ])', 'город ', text)
        text = re.sub(r'\bг\.\s*(?=[А-ЯЁ])', 'город ', text)

        # 3. АВТОМАТИЧЕСКОЕ СКЛОНЕНИЕ ВЕЛИЧИН И ЕДИНИЦ ИЗМЕРЕНИЯ
        units = {
            r'(\d+)\s*кг\b\.?': ["килограмм", "килограмма", "килограммов"],
            r'(\d+)\s*г\b\.?': ["грамм", "грамма", "граммов"],
            r'(\d+)\s*км\b\.?': ["километр", "километра", "километров"],
            r'(\d+)\s*м\b\.?': ["метр", "метра", "метров"],
            r'(\d+)\s*см\b\.?': ["сантиметр", "сантиметра", "сантиметров"],
            r'(\d+)\s*руб\b\.?': ["рубль", "рубля", "рублей"],
            r'(\d+)\s*р\b\.?': ["рубль", "рубля", "рублей"],
            r'(\d+)\s*коп\b\.?': ["копейка", "копейки", "копеек"],
        }

        for pattern, forms in units.items():
            def repl_unit(match, f=forms):
                num_str = match.group(1)
                chosen_form = get_declension(num_str, f)
                return f"{num_str} {chosen_form}"
            text = re.sub(pattern, repl_unit, text)
        # Фикс затыков на дефисах: заменяем обычный дефис между буквами на неразрывный (\u2011)
        text = re.sub(r'(?<=[а-яА-ЯёЁ])-(?=[а-яА-ЯёЁ])', '\u2011', text)
            
        return text

    def load_demagog_dictionaries(self):
        """Всеядный парсер v1.8.5: читает файлы баз в ОЗУ"""
        import os
        import glob

        dicts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "dicts")
        if not os.path.exists(dicts_dir):
            print("[СЛОВАРИ]: Папка 'dicts' не найдена.")
            return

        dic_files = glob.glob(os.path.join(dicts_dir, "*.dic"))
        for file_path in dic_files:
            file_name = os.path.basename(file_path)
            dict_key = os.path.splitext(file_name)[0]
            
            if dict_key not in self.dict_states:
                self.dict_states[dict_key] = False

            if "edge_Yo_" in dict_key:
                encodings_to_try = ["utf-8-sig", "utf-8", "cp1251"]
            else:
                encodings_to_try = ["utf-16", "utf-16-le", "utf-8-sig", "cp1251"]

            content = None
            for enc in encodings_to_try:
                try:
                    with open(file_path, "r", encoding=enc) as f:
                        content = f.read()
                        if content and len(content.strip()) > 0:
                            break
                except:
                    continue

            if not content:
                continue

            # Создаем плоскую изолированную структуру для каждого файла отдельно
            self.all_dicts_database[dict_key] = {"exact": [], "regex": []}

            for line in content.splitlines():
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                
                try:
                    left, right = line.split("=", 1)
                    left = left.strip()
                    right = right.strip()
                    
                    if not left or not right:
                        continue

                    if "*" in left:
                        self.all_dicts_database[dict_key]["regex"].append((left, right))
                    else:
                        self.all_dicts_database[dict_key]["exact"].append((left, right))
                except:
                    continue

        print(f"[СЛОВАРИ ИНИЦИАЛИЗАЦИЯ]: Базы прочитаны в ОЗУ. По дефолту все модули замен отключены.")
        self.rebuild_active_replacements()

    def rebuild_active_replacements(self):
        """Метод вызывается при переключении любого чекбокса. Безопасно собирает активный щит замен."""
        import re
        
        # Жестко очищаем старые рабочие кэши перед пересборкой
        self.exact_replacements = {}
        self.regex_replacements = []

        total_exact = 0
        total_regex = 0

        # Собираем заново только из тех файлов, где чекбокс = True
        for dict_key, is_active in self.dict_states.items():
            if is_active and dict_key in self.all_dicts_database:
                # 1. Переносим точные правила (ключ всегда в нижний регистр для регистронезависимости)
                for left, right in self.all_dicts_database[dict_key]["exact"]:
                    self.exact_replacements[str(left).lower()] = right
                    total_exact += 1
                
                # 2. Компилируем и переносим маски (*)
                for left, right in self.all_dicts_database[dict_key]["regex"]:
                    temp_left = left.replace("*", "___STAR___")
                    pattern = re.escape(temp_left).replace("___STAR___", r"[\w]*")
                    if not left.startswith("*"): 
                        pattern = r"\b" + pattern
                    if not left.endswith("*"): 
                        pattern = pattern + r"\b"
                    try:
                        self.regex_replacements.append((re.compile(pattern, re.IGNORECASE), right))
                        total_regex += 1
                    except:
                        continue

        print(f"[СЛОВАРИ АКТИВНЫ]: Собрано {total_exact} точных правил и {total_regex} масок на лету.")

    def _clean_text_notepad_killer(self, text):
        text = str(text)
        if not text.strip():
            return text

        # Нормализуем скрытые дефисы и символы тире
        text = text.replace("\u2011", "-").replace("‑", "-").replace("–", "-").replace("—", "-")

        # --- МОДУЛЬ СЛОВАРЕЙ DEMAGOG v1.8.5 ---
        # 1. Сначала прогоняем регулярные выражения со звёздочками (*)
        if hasattr(self, 'regex_replacements') and self.regex_replacements:
            for pattern, replacement in self.regex_replacements:
                text = pattern.sub(replacement, text)

        # 2. Прямая точечная замена по хэш-словарю с ручным контролем границ слова
        if hasattr(self, 'exact_replacements') and self.exact_replacements:
            sorted_keys = sorted(self.exact_replacements.keys(), key=len, reverse=True)
            for key in sorted_keys:
                lower_text = text.lower()
                start_pos = 0
                while True:
                    start_pos = lower_text.find(key, start_pos)
                    if start_pos == -1:
                        break
                    end_pos = start_pos + len(key)
                    is_start_bound = (start_pos == 0 or not text[start_pos - 1].isalnum())
                    is_end_bound = (end_pos == len(text) or not text[end_pos].isalnum())
                    
                    if is_start_bound and is_end_bound:
                        replacement = self.exact_replacements[key]
                        orig_word = text[start_pos:end_pos]
                        if orig_word.isupper() and len(orig_word) > 1:
                            final_rep = replacement.upper()
                        elif orig_word and orig_word.isupper():
                            final_rep = replacement.capitalize()
                        else:
                            final_rep = replacement
                        text = text[:start_pos] + final_rep + text[end_pos:]
                        lower_text = text.lower()
                        end_pos = start_pos + len(final_rep)
                    start_pos = end_pos
        # -------------------------------------

        # --- РОДНОЙ ВСТРОЕННЫЙ СЛОВАРЬ ESR ---
        if hasattr(self, 'user_dict') and self.user_dict:
            for raw_w, rep_w in self.user_dict.items():
                if raw_w.strip() and rep_w.strip():
                    lower_text = text.lower()
                    k = raw_w.strip().lower()
                    idx = 0
                    while True:
                        idx = lower_text.find(k, idx)
                        if idx == -1: break
                        e_idx = idx + len(k)
                        if (idx == 0 or not text[idx-1].isalnum()) and (e_idx == len(text) or not text[e_idx].isalnum()):
                            text = text[:idx] + rep_w + text[e_idx:]
                            lower_text = text.lower()
                            e_idx = idx + len(rep_w)
                        idx = e_idx
        # -------------------------------------

        # --- ФИНАЛЬНЫЙ СУПЕР-НОРМАЛИЗАТОР ЧИСЛИТЕЛЬНЫХ ДЛЯ ДВИЖКОВ TTS v1.8.5 ---
        # Этот блок перехватывает кривые компьютерные склейки словарей Demagog 
        # и превращает их в чистое, красивое человеческое произношение
        import re
        text = re.sub(r'\b10[-—‑\s]*ом\s+веке\b', 'десятом веке', text, flags=re.IGNORECASE)
        text = re.sub(r'\b10[-—‑\s]*го\s+века\b', 'десятого века', text, flags=re.IGNORECASE)
        text = re.sub(r'\b10[-—‑\s]*ый\s+век\b', 'десятый век', text, flags=re.IGNORECASE)
        text = re.sub(r'\b10\s+веке\b', 'десятом веке', text, flags=re.IGNORECASE)
        text = re.sub(r'\b10\s+век\b', 'десятый век', text, flags=re.IGNORECASE)
        text = re.sub(r'\b10\s+века\b', 'десятого века', text, flags=re.IGNORECASE)
        # ------------------------------------------------------------------------

        # Оригинальная очистка мусора
        text = text.replace("\r\n", "\n")
        text = text.replace("\r", "\n")
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            line_str = line.strip()
            if not line_str or line_str.startswith(("//", ";", "http")):
                continue
            cleaned_lines.append(line_str)
            
        return "\n".join(cleaned_lines)

    def _parse_path_safely(self, raw_path):
        from pathlib import Path
        if not raw_path: return ""
        if isinstance(raw_path, (tuple, list)) and len(raw_path) > 0: raw_path = raw_path
        p = Path(str(raw_path).replace("('", "").replace("',)", "").replace("',", "").replace("'", "").strip())
        return str(p.absolute())

    def _read_file_content(self, path):
        from pathlib import Path
        path_str = self._parse_path_safely(path)
        if not path_str or not os.path.exists(path_str): return ""
        ext = Path(path_str).suffix.lower()
        content = ""
        try:
            if ext == ".txt":
                for enc in ["utf-8-sig", "utf-8", "windows-1251", "cp1251"]:
                    try:
                        with open(path_str, "r", encoding=enc) as f: content = f.read()
                        if content: break
                    except: continue
            elif ext == ".fb2":
                for enc in ["utf-8-sig", "utf-8", "windows-1251"]:
                    try:
                        with open(path_str, "r", encoding=enc) as f: 
                            soup = BeautifulSoup(f.read(), "xml")
                        content = "\n".join([p.get_text() for p in soup.find_all("p")])
                        if content: break
                    except: continue
            elif ext == ".epub":
                import ebooklib
                from ebooklib import epub
                import logging
                logging.getLogger('ebooklib').setLevel(logging.ERROR)
                book = epub.read_epub(path_str)
                epub_paragraphs = []
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        soup = BeautifulSoup(item.get_content(), "html.parser")
                        for p in soup.find_all("p"):
                            txt = p.get_text().strip()
                            if txt: epub_paragraphs.append(txt)
                content = "\n".join(epub_paragraphs)
            elif ext == ".docx":
                import docx
                doc = docx.Document(path_str)
                docx_paragraphs = []
                for p in doc.paragraphs:
                    txt = p.text.strip()
                    if txt: docx_paragraphs.append(txt)
                content = "\n".join(docx_paragraphs)
        except Exception as e:
            print(f"[ERROR] Ошибка парсинга файла: {e}")
        return content

    def load_file(self):
        # --- ОЧИСТКА СТAРОГО КЭШA КОПИПAСТA ПРИ ЗАГРУЗКЕ КНИГИ ---
        try:
            if os.path.exists("esr_copy_cache.txt"):
                os.remove("esr_copy_cache.txt")
        except:
            pass
        # ---------------------------------------------------------

        if self.is_exporting: return
        file_path = filedialog.askopenfilename(filetypes=[
            ("Все поддерживаемые книги", "*.txt;*.fb2;*.epub;*.docx"),
            ("Текст (.txt)", "*.txt"),
            ("Книга (.fb2)", "*.fb2"),
            ("Книга (.epub)", "*.epub"),
            ("Документ Word (.docx)", "*.docx")
        ])
        if not file_path: return
        self.last_file_path = self._parse_path_safely(file_path)
        content = self._read_file_content(self.last_file_path)
        if content:
            normalized_content = self._normalize_text_dictionaries(content)
            cleaned_content = self._clean_text_notepad_killer(normalized_content)
            self.lbl_file.configure(text=os.path.basename(self.last_file_path), text_color="#1f8737")
            self.text_area.delete("1.0", "end")
            self.text_area.insert("1.0", cleaned_content)
            self.paragraphs, self.virtual_chunks = self._split_text_into_smart_chunks(cleaned_content)
            self.current_index = 0
            self.progress_bar.set(0)
            self.progress_label.configure(text="Прогресс: 0%")
            self.save_state()
            self.update_idletasks()

    # 📦 Блок 3: Стриминг-Плеер и Асинхронный Префетчер
    # Что внутри: Метод запуска потоков start_speech, фоновый префетчер _background_prefetch_loop и главный игровой цикл плеера _playlist_loop.
    # Что изменилось: Из префетчера и плеера убрана агрессивная перезапись файлов. Новая скорость подхватывается префетчером естественным путем — строго для тех чанков, которых еще нет на диске. Плеер больше не перегенерирует текущий кусок, исключая сбои. (05.07.2026 02:30)
    def start_speech(self):
        if self.is_exporting: return
        if self.is_paused:
            self.toggle_pause()
            return
        if self.is_speaking: return
        
        raw_text = self.text_area.get("1.0", "end").strip()
        if not raw_text or raw_text.startswith("Вставьте текст"): return
        
        # Нарезаем текст по новой виртуальной двухмерной схеме
        self.paragraphs, self.virtual_chunks = self._split_text_into_smart_chunks(raw_text)
        if not self.virtual_chunks: return
        
        self.is_speaking = True
        self.is_paused = False
        
        self.is_prefetching = True
        self.prefetch_thread = threading.Thread(target=self._background_prefetch_loop, daemon=True)
        self.prefetch_thread.start()
        
        threading.Thread(target=self._playlist_loop, args=(self,), daemon=True).start()

    def _get_formatted_speed(self):
        """v1.8.5-beta: Переводит шаги от -10 до 10 в безопасный формат % для Edge-TTS"""
        try:
            # Считаем процент отклонения: шаг +3 -> +30%
            step = int(self.speed)
            
            # === ЖЕСТКИЙ ЛИМИТЕР ДЛЯ ОНЛАЙНА ===
            # Edge-TTS часто сбоит, если темп выше +50% или ниже -50% при асинхронном стриминге
            if step > 5:
                step = 5
            elif step < -5:
                step = -5
                
            percent_offset = step * 10
            
            if percent_offset >= 0:
                return f"+{percent_offset}%"
            else:
                return f"{percent_offset}%"
        except:
            return "+0%"

    def _get_sapi5_speed(self):
        """v1.8.5-beta: Переводит шаги нашего ползунка (-10..10) в системный темп SAPI5 (-10..10)"""
        try:
            # SAPI5 принимает чистые целые числа от -10 до 10 напрямую. Нам даже конвертировать не нужно!
            step = int(self.speed)
            return max(-10, min(10, step))
        except:
            return 0

    def _get_formatted_volume(self):
        """Приводит self.volume к строгому стандарту edge-tts (например, '+0%', '-10%') или возвращает дефолт"""
        try:
            vol_str = str(self.volume).strip()
            # Если громкость дефолтная, edge-tts ожидает "+0%" или "+0Hz" для разных параметров. 
            # Безопаснее всего для стандартной громкости передавать "+0%"
            if vol_str == "100%" or vol_str == "100" or vol_str == "0%" or vol_str == "0" or vol_str == "+0%":
                return "+0%"
            
            if vol_str.startswith("+") or vol_str.startswith("-"):
                return vol_str if "%" in vol_str else f"{vol_str}%"
                
            if vol_str.isdigit():
                return f"+{vol_str}%"
        except:
            pass
        return "+0%"

    # --- УМНЫЙ ГЛУБОКИЙ ПРЕФЕТЧЕР В ОЗУ (ВЕРСИЯ 1.7.0) ---
    async def _download_chunk_to_ram(self, text, target_index):
        """Чистый асинхронный метод для скачивания чанка напрямую в ОЗУ"""
        try:
            safe_speed = self._get_formatted_speed()
            safe_volume = self._get_formatted_volume()
            communicate = edge_tts.Communicate(text, self.voice, rate=safe_speed, volume=safe_volume)
            
            audio_data = b""
            async for chunk in communicate.stream():
                if chunk["type"] == "audio":
                    audio_data += chunk["data"]
            
            if audio_data:
                self.ram_cache[target_index] = io.BytesIO(audio_data)
                print(f"[УСПЕХ RAM] Чанк {target_index} в ОЗУ. Размер: {len(audio_data)} байт.")
        except Exception as e:
            print(f"[ОШИБКА СЕРВЕРА] Чанк {target_index} не загружен: {e}")

    def _generate_audio_bytes_core(self, text, voice_name, loop=None):
        text = str(text).strip()
        if not text: return b""
        
        # Временная изоляция дефисов для защиты в регулярках словарей
        text = re.sub(r'(?<=[а-яА-ЯёЁ])-(?=[а-яА-ЯёЁ])', '\u2011', text)
        
        # --- ЛЕНИВАЯ ПРЕДОБРАБОТКА МИКРО-ЧАНКА В ПРЕФЕТЧЕРЕ (0.001 сек) ---
        if hasattr(self, '_normalize_text_dictionaries'):
            text = self._normalize_text_dictionaries(text)
        if hasattr(self, '_clean_text_notepad_killer'):
            text = self._clean_text_notepad_killer(text)
            
        # Возвращаем дефисы обратно перед отправкой в движки озвучки
        text = text.replace('\u2011', '-')
        
        target_id = self.voice_map.get(self.voice, self.voice)
        is_offline_mode = ("Офлайн" in str(self.voice) or "NATIVE" in str(target_id))
        
        # --- КРИТИЧЕСКИЙ ФИКС ДЛЯ ОНЛАЙН ДВИЖКА EDGE-TTS ---
        if not is_offline_mode:
            # Вычищаем спецсимволы словарей Demagog, которые вешают сервера Microsoft:
            # 1. Удаляем знаки "+" перед гласными (маркеры ударений Demagog)
            text = re.sub(r'\+([а-яА-ЯёЁё])', r'\1', text)
            # 2. Удаляем изолированные технические символы, знаки ударений (комбинируемые символы) и мусор
            text = text.replace("`", "").replace("'", "").replace("^", "").replace("_", "")
            # 3. Убираем двойные пробелы, которые могли возникнуть после замен
            text = re.sub(r'\s+', ' ', text).strip()
            
            # Если после очистки чанк стал пустым или в нем нет букв/цифр — блокируем отправку
            if not text or not re.search(r'[\w\d]', text):
                print("[ESR ЗАЩИТА]: Пропущен пустой/мусорный чанк для Edge-TTS.")
                return b""
        
        import win32com.client
        
        if is_offline_mode:
            try:
                speaker = win32com.client.Dispatch("SAPI.SpVoice")
                mem_stream = win32com.client.Dispatch("SAPI.SpMemoryStream")
                mem_stream.Format.Type = 34
                speaker.AudioOutputStream = mem_stream
                speaker.Rate = int(self.speed)
                
                voices = speaker.GetVoices()
                for idx in range(voices.Count):
                    if target_id in voices.Item(idx).Id:
                        speaker.Voice = voices.Item(idx)
                        break
                speaker.Speak(text)
                audio_bytes = bytes(mem_stream.GetData())
                return audio_bytes if audio_bytes else b""
            except Exception as e:
                print(f"[ОШИБКА SAPI5]: {e}")
                return b""
        else:
            try:
                import edge_tts
                import time
                
                # Микро-задержка 0.15 сек для защиты от Rate Limit
                time.sleep(0.15)
                
                safe_speed = f"{int(self.speed):+}%" if isinstance(self.speed, (int, float)) else "+0%"
                communicate = edge_tts.Communicate(text, target_id, rate=safe_speed, volume="+0%")
                
                audio_data = b""
                async def fetch_stream():
                    nonlocal audio_data
                    async_stream = communicate.stream()
                    async for chunk in async_stream:
                        if chunk["type"] == "audio":
                            audio_data += chunk["data"]
                if loop:
                    loop.run_until_complete(fetch_stream())
                
                # Если сервер вернул пустоту, принудительно вызываем ошибку для ухода в Fallback
                if not audio_data:
                    raise RuntimeWarning("Сервер Edge вернул пустой аудио-поток.")
                    
                return audio_data
                
            except Exception as e:
                print(f"[Резерв ESR]: Онлайн-движок споткнулся ({e}). Переключаем чанк на Офлайн SAPI5...")
                # --- АВАРИЙНЫЙ КАНАЛ (FALLBACK): Генерируем этот же чанк через локальный SAPI5 ---
                try:
                    import win32com.client
                    speaker = win32com.client.Dispatch("SAPI.SpVoice")
                    mem_stream = win32com.client.Dispatch("SAPI.SpMemoryStream")
                    mem_stream.Format.Type = 34  # WAV PCM 22050Hz 16-bit Mono
                    speaker.AudioOutputStream = mem_stream
                    
                    # === ИСПРАВЛЕНИЕ v1.8.5-beta: Безопасный темп SAPI5 ===
                    speaker.Rate = self._get_sapi5_speed()
                    
                    speaker.Speak(text)
                    audio_bytes = bytes(mem_stream.GetData())
                    if audio_bytes:
                        print("[Резерв ESR УСПЕХ]: Чанк успешно спасен локальным движком.")
                        # Возвращаем кортеж (байты, флаг_формата), чтобы плеер знал, что это WAV
                        return (audio_bytes, "WAV")
                except Exception as sapi_err:
                    print(f"[КРИТИЧЕСКАЯ ОШИБКА РЕЗЕРВА]: Не удалось озвучить даже через SAPI5: {sapi_err}")
                
                return (b"", "MP3")

    def _background_prefetch_loop(self):
        """
        Параллельный асинхронный префетчер ОЗУ (Глубина буфера: 4 чанка вперед).
        Анализирует маркер чтения, на лету запрашивает аудио у серверов Microsoft 
        или генерирует через SAPI5, упаковывая данные в потоки io.BytesIO.
        Устраняет паузы между абзацами при воспроизведении.
        """
        import time
        import asyncio
        import io
        
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        buffer_depth = 4  # Загружаем на 4 чанка вперед

        while self.is_speaking and self.is_prefetching:
            # Находим список индексов, которые нужно срочно загрузить
            targets_to_load = []
            for step in range(0, buffer_depth + 1):
                check_idx = self.current_chunk_index + step
                if check_idx >= len(self.virtual_chunks):
                    break
                if check_idx not in self.ram_cache:
                    targets_to_load.append(check_idx)

            if targets_to_load:
                target_id = self.voice_map.get(self.voice, self.voice)
                is_offline = "Офлайн" in str(self.voice) or "NATIVE" in str(self.voice) or "NATIVE" in str(target_id)

                for next_target in targets_to_load:
                    # Проверка: не ушел ли маркер вперед, пока мы в цикле
                    if next_target < self.current_chunk_index:
                        continue

                    chunk_data = self.virtual_chunks[next_target]
                    chunk_text = chunk_data[1] if isinstance(chunk_data, tuple) and len(chunk_data) > 1 else chunk_data
                    
                    if not str(chunk_text).strip():
                        continue

                    p_text = self.optimize_text_pauses(str(chunk_text))
                    
                    # Синтезируем аудио (онлайн или офлайн)
                    audio_bytes = self._generate_audio_bytes_core(p_text, self.voice, loop=loop)
                    
                    # Проверяем актуальность маркера после скачивания
                    if audio_bytes and len(audio_bytes) > 0 and next_target >= self.current_chunk_index:
                        # Записываем строго валидный, неповрежденный поток аудио
                        self.ram_cache[next_target] = io.BytesIO(audio_bytes)
                        print(f"[УСПЕХ RAM v1.8.5] Чанк {next_target} успешно загружен в ОЗУ. ({len(audio_bytes)} байт)")
                
                # Микро-пауза цикла для разгрузки процессора
                time.sleep(0.02)
            else:
                time.sleep(0.1)

        loop.close()

    # --- ОСНОВНОЙ ПЛЕЙЛИСТ-ЦИКЛ ВОСПРОИЗВЕДЕНИЯ ИЗ ОЗУ (ВЕРСИЯ 1.8.5 - FIX PROGRESS) ---
    def _playlist_loop(self, *args):
        """
        Главный управляющий поток гибридного плеера (ОЗУ-ориентированный).
        Непрерывно опрашивает mixer, извлекает байты из self.ram_cache[idx], 
        автоматически переключает частоту звука под онлайн (MP3) или офлайн (PCM/WAV) движки
        и координирует оранжевую подсветку текста в интерфейсе.
        """
        import time
        import pygame
        import io

        while self.is_speaking:
            if self.is_paused:
                time.sleep(0.1)
                continue

            if not pygame.mixer.music.get_busy() and not pygame.mixer.get_busy():
                idx = self.current_chunk_index
                total_chunks = len(self.virtual_chunks) if hasattr(self, 'virtual_chunks') else 0
                
                if idx < total_chunks:
                    if idx in self.ram_cache:
                        try:
                            bytes_stream = self.ram_cache[idx]
                            bytes_stream.seek(0)
                            raw_bytes = bytes_stream.getvalue()

                            if len(raw_bytes) > 0:
                                # --- СИНХРОНИЗАЦИЯ ПОДСВЕТКИ С ЯДРОМ 1.7.0 ---
                                chunk_data = self.virtual_chunks[idx]
                                if isinstance(chunk_data, tuple) and len(chunk_data) > 0:
                                    actual_para_idx = chunk_data[0]
                                    self.after(0, lambda p=actual_para_idx: self._highlight_paragraph_on_screen(p))
                                # --------------------------------------------

                                # --- ЖЕСТКИЙ ВОЗВРАТ АНИМАЦИИ ПРОГРЕСС-БАРА ЧТЕНИЯ ---
                                if total_chunks > 0:
                                    progress_float = idx / total_chunks
                                    self.after(0, lambda p=progress_float, c=idx+1, t=total_chunks: 
                                        self._update_ui_progress(p, f"Чтение чанков: {int(p*100)}% ({c}/{t})"))
                                # ----------------------------------------------------

                                # --- ГИБРИДНЫЙ АУДИОПЛЕЕР С ПОДДЕРЖКОЙ PCM И MP3 ИЗ ОЗУ ---
                                target_id = self.voice_map.get(self.voice, self.voice) if hasattr(self, 'voice_map') else self.voice
                                is_offline = "Офлайн" in str(self.voice) or "NATIVE" in str(self.voice) or "NATIVE" in str(target_id)
                                
                                # === АКТУАЛИЗАЦИЯ v1.8.5-beta: Умный динамический реинит частоты ===
                                is_offline = "Офлайн" in str(self.voice) or "NATIVE" in str(self.voice) or "NATIVE" in str(target_id)
                                
                                # Определяем целевую частоту: SAPI5 выдает 22050 Гц, онлайн Edge-TTS выдает 24000 Гц
                                target_freq = 22050 if is_offline else 24000
                                
                                # Проверяем, совпадает ли текущая частота микшера с целевой
                                current_mixer_freq = pygame.mixer.get_init()[0] if pygame.mixer.get_init() else None
                                
                                if current_mixer_freq != target_freq:
                                    try:
                                        if pygame.mixer.music.get_busy():
                                            pygame.mixer.music.stop()
                                        pygame.mixer.quit()
                                        
                                        # Инициализируем микшер под конкретный тип чанка
                                        if is_offline:
                                            pygame.mixer.init(frequency=22050, size=-16, channels=1)
                                        else:
                                            pygame.mixer.init(frequency=24000, size=-16, channels=2)
                                    except Exception as init_err:
                                        print(f"[ОШИБКА РЕИНИТА МИКСЕРА]: {init_err}")

                                # Запускаем воспроизведение
                                try:
                                    if is_offline:
                                        sound_obj = pygame.mixer.Sound(buffer=raw_bytes)
                                        sound_obj.play()
                                        
                                        # Ожидаем окончания проигрывания офлайн-звука
                                        while pygame.mixer.get_busy() and self.is_speaking and not self.is_paused:
                                            time.sleep(0.05)
                                    else:
                                        play_stream = io.BytesIO(raw_bytes)
                                        play_stream.seek(0)
                                        pygame.mixer.music.load(play_stream)
                                        pygame.mixer.music.play()
                                except Exception as mix_err:
                                    print(f"[СБОЙ МИКСЕРА]: {mix_err}")
                                # ------------------------------------------------------

                                if not self.is_paused:
                                    self.current_chunk_index += 1
                                    # Если это офлайн (WAV/PCM), оставляем микропаузу для безопасности, онлайн гоним без задержек
                                    if is_offline:
                                        time.sleep(0.03)
                        except Exception as e:
                            print(f"[ОШИБКА ЦИКЛА ПЛЕЕРА]: {e}")
                            time.sleep(0.05)
                else:
                    # Если дочитали до конца книги — плавно сбрасываем UI в 100%
                    self.after(0, lambda: self._update_ui_progress(1.0, "Завершено! 100%"))
                    time.sleep(0.1)

    # 📦 Блок 4: Подсветка текста, Управление кнопками и Экспорт MP3
    # Что внутри: Функции _highlight_paragraph_on_screen, toggle_pause, stop_speech, деструктор on_close и поабзацный метод сохранения export_mp3.
    def _highlight_paragraph_on_screen(self, para_idx):
        if not self.is_speaking or para_idx < 0 or para_idx >= len(self.paragraphs): return
        self.text_area.tag_remove("para_highlight", "1.0", "end")
        
        # Достаем текст оригинального крупного абзаца
        paragraph_text = self.paragraphs[para_idx]
        if not paragraph_text: return
        
        idx = self.text_area._textbox.search(paragraph_text, "1.0", stopindex="end")
        if idx:
            end_idx = f"{idx} + {len(paragraph_text)} chars"
            self.text_area.tag_add("para_highlight", idx, end_idx)
            self.text_area.see(idx)

    # --- УМНАЯ СИСТЕМА ВОЗОБНОВЛЕНИЯ ПОТОКОВ ЧТЕНИЯ ---
    def toggle_pause(self):
        if self.is_exporting: return
        
        if self.is_paused and not pygame.mixer.music.get_busy() and self.prefetch_thread is None:
            self.is_speaking = True
            self.is_paused = False
            self.is_prefetching = True
            
            self.after(0, lambda: self.btn_pause.configure(text="⏸  Пауза", fg_color="#d97706", hover_color="#b45309"))
            
            self.prefetch_thread = threading.Thread(target=self._background_prefetch_loop, daemon=True)
            self.prefetch_thread.start()
            threading.Thread(target=self._playlist_loop, args=(self,), daemon=True).start()
            return

        if not self.is_speaking: return
        
        if not self.is_paused:
            self.is_paused = True
            pygame.mixer.music.pause()
            self.save_state()
            self.after(0, lambda: self.btn_pause.configure(text="►  Продолжить", fg_color="#10b981", hover_color="#059669"))
        else:
            self.is_paused = False
            pygame.mixer.music.unpause()
            self.after(0, lambda: self.btn_pause.configure(text="⏸  Пауза", fg_color="#d97706", hover_color="#b45309"))

    def stop_speech(self, clear_file_path=True):
        if self.is_exporting: return
        self.is_speaking = False
        self.is_paused = False
        self.is_prefetching = False
        self.prefetch_thread = None
        
        if clear_file_path:
            self.last_file_path = ""
            self.lbl_file.configure(text="Файл не выбран", text_color="gray")
            self.current_chunk_index = 0
            self.progress_bar.set(0)
            self.progress_label.configure(text="Прогресс: 0%")
            
            try:
                self.text_area.delete("1.0", "end")
            except:
                pass
                
            # --- ХАРД-ОЧИСТКА ВНЕШНЕГО КЭША ПРИ НАЖАТИИ СТОП ---
            try:
                cache_file = "esr_copy_cache.txt"
                if os.path.exists(cache_file):
                    os.remove(cache_file)
            except:
                pass
            # --------------------------------------------------
            
        self.save_state()
        self.after(0, lambda: self.btn_pause.configure(text="⏸  Пауза", fg_color="#d97706", hover_color="#b45309"))
        self.text_area.tag_remove("para_highlight", "1.0", "end")
        
        try:
            pygame.mixer.music.stop()
            pygame.mixer.music.unload()
            pygame.mixer.stop() # Жестко глушит все Sound каналы (WAV)
        except:
            pass
        
        # Полное освобождение оперативной памяти от звуковых чанков
        if hasattr(self, 'ram_cache'):
            self.ram_cache.clear()

    def on_close(self):
        self.is_speaking = False
        self.is_prefetching = False
        try:
            pygame.mixer.music.stop()
            pygame.mixer.music.unload()
        except:
            pass
            
        # Защитная очистка ОЗУ перед выходом
        if hasattr(self, 'ram_cache'):
            self.ram_cache.clear()
            
        self.save_state_explicit(clear_path=False)
        self.destroy()

    def export_mp3(self):
        if (self.is_speaking and not self.is_paused) or self.is_exporting: 
            return
        
        has_selection = False
        try:
            if self.text_area._textbox.tag_ranges("sel"):
                has_selection = True
        except:
            pass

        if has_selection:
            selected_text = self.text_area.get("sel.first", "sel.last").strip()
            if not selected_text: return
            cleaned_text = self._clean_text_notepad_killer(selected_text)
            export_paragraphs = [p.strip() for p in cleaned_text.split("\n") if p.strip()]
        else:
            if not self.virtual_chunks: 
                messagebox.showinfo("Инфо", "Нет загруженной книги для экспорта!")
                return
            
            export_paragraphs = []
            start_idx = getattr(self, "current_chunk_index", 0)
            if start_idx >= len(self.virtual_chunks):
                start_idx = 0

            for idx in range(start_idx, len(self.virtual_chunks)):
                chunk_data = self.virtual_chunks[idx]
                if isinstance(chunk_data, (tuple, list)) and len(chunk_data) > 1:
                    chunk_text = chunk_data[1]
                else:
                    chunk_text = chunk_data
                
                if str(chunk_text).strip():
                    export_paragraphs.append(str(chunk_text).strip())

        if not export_paragraphs: 
            messagebox.showinfo("Инфо", "Текст для сохранения не найден!")
            return

        target_id = self.voice_map.get(self.voice, self.voice)
        is_offline = "Офлайн" in str(self.voice) or "NATIVE" in str(self.voice) or "NATIVE" in str(target_id)
        
        default_ext = ".wav" if is_offline else ".mp3"
        file_types = [("Аудиофайл WAV", "*.wav")] if is_offline else [("Аудиофайл MP3", "*.mp3")]
        
        save_path = filedialog.asksaveasfilename(defaultextension=default_ext, filetypes=file_types)
        if not save_path: return

        self.is_exporting = True
        self.btn_export.configure(state="disabled", text="⚡ Экспорт...")

        def create_wav_header(pcm_data_size, sample_rate=24000, bits_per_sample=16, channels=2):
            import struct
            num_channels = channels
            bytes_per_sample = bits_per_sample // 8
            byte_rate = sample_rate * num_channels * bytes_per_sample
            block_align = num_channels * bytes_per_sample
            return struct.pack('<4sI4s4sIHHIIHH4sI',
                b'RIFF', 36 + pcm_data_size, b'WAVE', b'fmt ',
                16, 1, num_channels, sample_rate, byte_rate, block_align, bits_per_sample,
                b'data', pcm_data_size
            )

        def run_export():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            total_p = len(export_paragraphs)
            ram_audio_chunks = []
            
            try:
                if is_offline:
                    for idx, para in enumerate(export_paragraphs):
                        p_text = self.optimize_text_pauses(para)
                        audio_bytes = self._generate_audio_bytes_core(p_text, self.voice, loop=loop)
                        if audio_bytes and len(audio_bytes) > 0:
                            ram_audio_chunks.append(audio_bytes)
                        
                        prog = int(((idx + 1) / total_p) * 100)
                        # Обновляем текст прямо на кнопке экспорта, не трогая плеер
                        self.after(0, lambda p=prog: self.btn_export.configure(text=f"⚡ {p}%"))
                    
                    if ram_audio_chunks:
                        raw_pcm_monolith = b"".join(ram_audio_chunks)
                        with open(save_path, "wb") as outfile:
                            outfile.write(create_wav_header(len(raw_pcm_monolith)))
                            outfile.write(raw_pcm_monolith)
                else:
                    with open(save_path, "wb") as outfile:
                        for idx, para in enumerate(export_paragraphs):
                            p_text = self.optimize_text_pauses(para)
                            audio_bytes = self._generate_audio_bytes_core(p_text, self.voice, loop=loop)
                            if audio_bytes and len(audio_bytes) > 0:
                                outfile.write(audio_bytes)
                            
                            prog = int(((idx + 1) / total_p) * 100)
                            self.after(0, lambda p=prog: self.btn_export.configure(text=f"⚡ {p}%"))

                fmt_name = "WAV" if is_offline else "MP3"
                self.after(0, lambda f=fmt_name: messagebox.showinfo("Успех ESR", f"Аудио успешно сохранено в файл ({f})!"))
            
            except Exception as e:
                self.after(0, lambda err=e: messagebox.showerror("Ошибка", f"Не удалось экспортировать: {err}"))
            finally:
                ram_audio_chunks.clear()
                self.is_exporting = False
                # Возвращаем исходный текст кнопке
                self.after(0, lambda: self.btn_export.configure(state="normal", text=" 💾 Сохранить Аудио"))

        threading.Thread(target=run_export, daemon=True).start()

    # 📦 Блок 5: Сохранение Состояния JSON и Точка Входа (Конец файла)
    # What inside: Методы save_state, save_state_explicit, load_state и системный запуск if __name__ == "__main__":.
    # Что изменилось: В JSON-структуру добавлены поля "theme" и "text_bg". Метод load_state теперь полноценно восстанавливает тему и фоны на лету при старте приложения.(05.07.2026 01:10)
    def save_state(self):
        if not self.last_file_path:
            self.save_state_explicit(clear_path=True)
        else:
            self.save_state_explicit(clear_path=False)

    def save_state_explicit(self, clear_path=False):
        """Защищенное сохранение кэша копипаста без риска автоудаления v1.7.2"""
        try:
            cache_file = "esr_copy_cache.txt"
            raw_text = self.text_area.get("1.0", "end").strip()
            if raw_text.startswith("Вставьте текст"):
                raw_text = ""

            # Если файла книги нет, но текст в окне присутствует — пишем его в кэш
            if not self.last_file_path and raw_text and not clear_path:
                with open(cache_file, "w", encoding="utf-8") as cache_f:
                    cache_f.write(raw_text)
            elif self.last_file_path:
                # Если открыт официальный файл, кэш копипаста больше не нужен
                if os.path.exists(cache_file):
                    try: os.remove(cache_file)
                    except: pass

            state = {
                "last_file_path": "" if clear_path else self.last_file_path,
                "voice": self.voice,
                "speed": self.speed,
                "current_chunk_index": self.current_chunk_index,
                "theme": self.theme,
                "text_bg": self.text_bg,
                "lang": self.lang  # <-- НОВАЯ СТРОКА: сохраняем "RU" или "EN"
            }
            
            with open(self.config_file, "w", encoding="utf-8") as f:
                json.dump(state, f, ensure_ascii=False, indent=4)
        except Exception as e:
            print(f"[ОШИБКА СОХРАНЕНИЯ КОНФИГА]: {e}")

    def load_state(self):
        if not os.path.exists(self.config_file): return
        try:
            with open(self.config_file, "r", encoding="utf-8") as f: 
                state = json.load(f)
            
            # Читаем язык. Если файла настроек еще нет, по дефолту берем "RU"
            self.lang = state.get("lang", "RU")

            self.refresh_interface_texts()
            self.last_file_path = state.get("last_file_path", "")
            saved_voice = state.get("voice", "ru-RU-DmitryNeural")
            
            # --- ЗАЩИТА И АДАПТАЦИЯ СКОРОСТИ ПОД НОВУЮ ШКАЛУ ---
            saved_speed_raw = state.get("speed", 0)
            try:
                # Очищаем строку от процентов и знаков, если они там есть, и превращаем в число
                saved_speed = int(str(saved_speed_raw).replace("%", "").replace("+", ""))
            except:
                saved_speed = 0

            if saved_speed > 10 or saved_speed < -10:
                saved_speed = 0
            
            self.speed_slider.set(saved_speed)
            self.speed = saved_speed
            
            # Синхронизируем текстовое отображение плашки в GUI
            self.on_speed_change(saved_speed)
            
            self.current_chunk_index = state.get("current_chunk_index", 0)
            
            self.theme = state.get("theme", "Светлая (Light)")
            if "Тёмная" in self.theme: ctk.set_appearance_mode("Dark")
            else: ctk.set_appearance_mode("Light")
                
            self.text_bg = state.get("text_bg", "«Стандартный»")
            self.change_text_bg(self.text_bg)
            
            for k, v in self.voice_map.items():
                if v == saved_voice: self.voice = v; self.voice_combo.set(k); break
                
            # --- ВЕТКА 1: ЕСЛИ БЫЛ ОТКРЫТ ОФИЦИАЛЬНЫЙ ФАЙЛ КНИГИ ---
            if self.last_file_path:
                cleaned = self._parse_path_safely(self.last_file_path)
                if os.path.exists(cleaned):
                    content = self._read_file_content(cleaned)
                    if content:
                        self.last_file_path = cleaned
                        self.lbl_file.configure(text=os.path.basename(cleaned), text_color="#1f8737")
                        self.text_area.delete("1.0", "end")
                        self.text_area.insert("1.0", content)
                        
                        self.paragraphs, self.virtual_chunks = self._split_text_into_smart_chunks(content)
                        total_chunks = len(self.virtual_chunks)
                        
                        if total_chunks and self.current_chunk_index < total_chunks:
                            self.is_speaking = False
                            self.is_paused = True
                            self.btn_pause.configure(text="►  Продолжить", fg_color="#10b981", hover_color="#059669")
                            
                            prog = self.current_chunk_index / total_chunks
                            self._update_ui_progress(prog, f"Пауза: {int(prog*100)}% ({self.current_chunk_index}/{total_chunks})")
                            
                            chunk_data = self.virtual_chunks[self.current_chunk_index]
                            if isinstance(chunk_data, tuple) and len(chunk_data) > 0:
                                para_idx = chunk_data[0]
                                self._highlight_paragraph_on_screen(para_idx)
            
            # --- ВЕТКА 2: ТЕКСТОВЫЙ КЭШ КОПИПАСТА (ЕСЛИ ФАЙЛА НЕ БЫЛО) ---
            else:
                cache_file = "esr_copy_cache.txt"
                if os.path.exists(cache_file):
                    try:
                        with open(cache_file, "r", encoding="utf-8") as cache_f:
                            saved_text = cache_f.read().strip()
                        
                        if saved_text:
                            self.text_area.delete("1.0", "end")
                            self.text_area.insert("1.0", saved_text)
                            
                            self.paragraphs, self.virtual_chunks = self._split_text_into_smart_chunks(saved_text)
                            total_chunks = len(self.virtual_chunks)
                            
                            if total_chunks and self.current_chunk_index < total_chunks:
                                self.is_speaking = False
                                
                                # Принудительно взводим флаги паузы, чтобы активировать продолжение
                                self.is_paused = True
                                self.btn_pause.configure(text="►  Продолжить", fg_color="#10b981", hover_color="#059669")
                                
                                prog = self.current_chunk_index / total_chunks
                                self._update_ui_progress(prog, f"Пауза: {int(prog*100)}% ({self.current_chunk_index}/{total_chunks})")
                                
                                chunk_data = self.virtual_chunks[self.current_chunk_index]
                                if isinstance(chunk_data, tuple) and len(chunk_data) > 0:
                                    para_idx = chunk_data[0]
                                    self._highlight_paragraph_on_screen(para_idx)
                    except Exception as cache_err:
                        print(f"[ОШИБКА ЧТЕНИЯ КЭША КОПИПАСТА]: {cache_err}")
                        
        except Exception as e: 
            print(f"[ОШИБКА ЗАГРУЗКИ КОНФИГА]: {e}")

    def _update_ui_progress(self, val, text):
        self.progress_bar.set(val)
        self.progress_label.configure(text=text)

    # === МЕТОД ОТКРЫТИЯ ОКНА ЧТЕНИЯ (ВСТАВИТЬ В КЛАСС EdgeStreamReaderApp) ===
    def open_reader_window(self):
        """Открывает независимое окно для чтения текста"""
        # Защита от дублирования: если окно уже открыто, выводим его на передний план
        if hasattr(self, 'reader_win') and self.reader_win.winfo_exists():
            self.reader_win.lift()
            return
            
        # Создаем окно и передаем ему ссылку на наш главный текстовый блок
        self.reader_win = ESRReaderWindow(self, self.text_area)

    def _open_local_readme(self, filename):
        """v1.8.5: Безопасный запуск локального файла разметки в ОС Windows по клику из GUI"""
        import os
        if os.path.exists(filename):
            try:
                os.startfile(filename)
            except Exception as e:
                from tkinter import messagebox
                messagebox.showerror("Ошибка", f"Не удалось открыть файл: {e}")
        else:
            from tkinter import messagebox
            messagebox.showwarning("Внимание", f"Файл {filename} не найден в корневой директории программы.\nПожалуйста, создайте его.")

class ESRReaderWindow(ctk.CTkToplevel):
    """
    Независимый интерактивный экран чтения (функционал суфлёра-планшета).
    Синхронизирует позицию скролла с ядром, полностью подчиняется цветовой схеме
    главного окна и делегирует сохранение закладок через координатную карту инсерта.
    """
    def __init__(self, parent, main_text_area):
        # Конструктор v1.8.5 с защитой от сброса иконки Toplevel в стандартное перо Tkinter
        super().__init__(parent)
        self.parent = parent
        self.main_text_area = main_text_area
        
        self.title("ESR v1.8.5 — Экран чтения")
        self.geometry("900x700")
        self.minsize(500, 400)
        
        # Принудительно выводим окно поверх остальных
        self.title(parent.tr("reader_title"))

        # Настройка сетки: Строка 0 для панели, Строка 1 для текста
        self.grid_rowconfigure(0, weight=0)
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)
        
        # --- МИНИМАЛИСТИЧНАЯ ПАНЕЛЬ УПРАВЛЕНИЯ ---
        self.top_panel = ctk.CTkFrame(self, fg_color="transparent")
        self.top_panel.grid(row=0, column=0, sticky="ew", padx=10, pady=5)
        
        self.btn_add_bookmark = ctk.CTkButton(
            self.top_panel, 
            text=parent.tr("reader_bookmark"),
            command=self.set_bookmark_from_cursor,
            width=180
        )
        self.btn_add_bookmark.pack(side="left", padx=5)
        
        # --- ТЕКСТОВОЕ ПОЛЕ ДЛЯ ЧТЕНИЯ ---
        self.reader_text = ctk.CTkTextbox(
            self,
            font=("Segoe UI", 18),
            wrap="word",
            padx=25,
            pady=25
        )
        self.reader_text.grid(row=1, column=0, sticky="nsew")

        # Инициализация цвета при создании окна на основе текущих настроек ядра
        current_bg_color = parent.text_area.cget("fg_color")
        current_text_color = parent.text_area.cget("text_color")
        self.reader_text.configure(fg_color=current_bg_color, text_color=current_text_color)

        # Переносим text книги в окно чтения с проверкой дефолтной строки
        raw_text = self.main_text_area.get("1.0", "end").strip()

        if raw_text in ["Вставьте текст или откройте файл книги...", "Paste text or open a book file..."]:
            self.reader_text.insert("1.0", parent.tr("text_area_init"))
        else:
            self.reader_text.insert("1.0", self.main_text_area.get("1.0", "end"))
        
        # Синхронизируем положение при старте окна
        self.sync_with_active_chunk()

        # === ДОРАБОТКА v1.8.5: Защита иконки окна чтения от сброса в дефолт ===
        if os.path.exists("icon.ico"):
            try:
                self.wm_iconbitmap("icon.ico")
            except:
                pass
        else:
            # Если иконки на диске нет, убираем перо Tkinter через тактовую паузу
            self.after(200, lambda: self.iconbitmap(default=""))
        
    def set_bookmark_from_cursor(self):
        """Передает точный текст строки в главное ядро для безошибочного сохранения"""
        try:
            # 1. Получаем номер строки, на которой стоит курсор в окне чтения
            cursor_pos = self.reader_text.index("insert")
            line_num = cursor_pos.split('.')[0]
            
            # 2. Вытаскиваем чистый текст этого абзаца из окна чтения
            current_line_text = self.reader_text.get(f"{line_num}.0", f"{line_num}.end").strip()
            
            # Если строка пустая (кликнули между абзацами), берем строку чуть ниже
            if not current_line_text or len(current_line_text) < 5:
                line_num = str(int(line_num) + 1)
                current_line_text = self.reader_text.get(f"{line_num}.0", f"{line_num}.end").strip()

            if current_line_text:
                # 3. Сохраняем этот текст в специальную переменную главного класса
                self.parent._reader_selected_text = current_line_text
                
                # 4. Программно вызываем родной метод сохранения в JSON
                if hasattr(self.parent, 'add_bookmark'):
                    self.parent.add_bookmark()
                    
                print(f"[ESR ОКНО ЧТЕНИЯ]: Фраза передана в главное ядро для прямой привязки.")
            
        except Exception as e:
            print(f"[ESR ОКНО ЧТЕНИЯ ОШИБКА СОХРАНЕНИЯ ЗАКЛАДКИ]: {e}")
            
    def sync_with_active_chunk(self):
        """Считывает маркер из главного окна и принудительно скроллит экран чтения на старте"""
        try:
            main_cursor_pos = self.main_text_area.index("insert")
            self.reader_text.mark_set("insert", main_cursor_pos)
            self.reader_text.see(main_cursor_pos)
        except Exception as e:
            print(f"[ESR ОКНО ЧТЕНИЯ ОШИБКА ПЕРВИЧНОЙ СИНХРОНИЗАЦИИ]: {e}")

if __name__ == "__main__":
    app = EdgeStreamReaderApp()
    app.mainloop()